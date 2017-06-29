""" The simulation package for the list of plans and the plans class.
    Author: Shannon Grubb
            shannon.grubb@nist.gov
    2017-05
"""

import csv
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENTATION
from tkinter import filedialog

from Data.ClassBenefits import Benefits, Benefit
from Data.ClassCosts import Costs
from Data.ClassExternalities import Externalities
from Data.ClassFatalities import Fatalities
from Data.ClassNonDBens import NonDBens
from Data.ClassNonDBens import Benefit as NonDBenefit

from Data.distributions import uniDistInv, triDistInv, gauss_dist_inv, none_dist, discrete_dist_inv


from NewIRR import irr_for_all
from GUI.Constants import TAB

import math
import numpy as np

class Simulation():
    """ Holds all of the plans and does all of the larger calculations. """
    parties = ['developer', 'title holder(s)', 'lender(s)', 'tenants', 'users', 'community']
    def __init__(self):
        self.title = ""
        self.plan_list = []
        self.num_plans = 0
        self.horizon = 0
        self.discount_rate = 0.03
        self.risk_pref = "neutral"
        self.stat_life = 7500000

    def file_read(self, file_name):
        """ Reads in the save file to fill the simulation with all plans. """
        file_data = open(file_name, 'r')
        data = csv.reader(file_data)
        for line in data:
            if line[0] == 'Title':
                self.title = line[1]
                self.horizon = line[4]
            elif line[0] == 'Discount Rate':
                self.discount_rate = line[1]
                self.risk_pref = line[4]
                self.stat_life = line[7]
            elif 'Plan' in line[0]:
                build_list = [list(line)]
                next_plan = Plan(line[0][-1], line[1], [line[2], list(line[3:9])],
                                 [line[9], list(line[10:17])],
                                 self.discount_rate, self.horizon, self.stat_life, self.parties)
            elif line[0] == "":
                build_list.append(list(line))
            elif line[0] == 'END PLAN':
                next_plan.new_plan(build_list)
                self.plan_list.append(next_plan)
                build_list = []
            elif line[0] == 'END FILE':
                self.num_plans = len(self.plan_list)

                return True
            else:
                pass

    def file_save(self):
        """ Saves the data in the save file."""
        my_formats = [('Comma Separated Value', '*.csv'),]
        file_name = filedialog.asksaveasfilename(filetypes=my_formats, title="Save the file as...")
        if '.csv' != file_name[-4:]:
            file_name = file_name + '.csv'
        new_file = open(file_name, 'w')
        new_file.write('Title,' + str(self.title) + ',,Horizon,' + str(int(self.horizon)) + '\n')
        new_file.write('Discount Rate,' + str(self.discount_rate))
        new_file.write(',,Risk Prev,' + self.risk_pref)
        new_file.write(',,Value of Stat Life,' + str(self.stat_life) + '\n')
        new_file.write('BEGIN PLANS\n')
        for plan in self.plan_list:
            plan.save_plan(new_file)
            new_file.write('END PLAN\n')
        new_file.write('END FILE')
        new_file.close()

    def csv_export(self):
        my_formats = [('Comma Separated Value', '*.csv'),]
        file_name = filedialog.asksaveasfilename(filetypes=my_formats, title="Save the file as...")
        if '.csv' != file_name[-4:]:
            file_name = file_name + '.csv'
        new_file = open(file_name, 'w')
        new_file.write('Outputs of Economic Evaluation: [' + self.title + ']\n')
        new_file.write(',Base Case')
        for i in range(1, self.num_plans):
            new_file.write(',Alternative ' + str(i))
        new_file.write('\n,Base')
        for i in range(1, self.num_plans):
            new_file.write(',' + self.plan_list[i].name)
        new_file.write('\nBenefits\nDisaster Economic Benefits\nResponse and Recovery Costs')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.bens.r_sum))
        new_file.write('\nDirect Loss Reduction')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.bens.d_sum))
        new_file.write('\nIndirect Losses')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.bens.i_sum))
        new_file.write('\nDisaster Non-Market Benefits\nValue of Statistical Lives Saved')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.fat.stat_value_averted))
        new_file.write('\nNumber of Statistical Lives Saved')
        for plan in self.plan_list:
            new_file.write(',' + str(plan.fat.stat_averted))
        new_file.write('\nNon-disaster Related Benefits\nOne-Time')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.nond_bens.one_sum))
        new_file.write('\nRecurring\n')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.nond_bens.r_sum))
        new_file.write('\nCosts\nDirect Costs')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.costs.d_sum))
        new_file.write('\nIndirect Costs')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.costs.i_sum))
        new_file.write('\nOMR\nOne-Time')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.costs.omr_1_sum))
        new_file.write('\nRecurring')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.costs.omr_r_sum))
        new_file.write('\nExternalities\nPositive\nOne-Time')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.exts.one_sum_p))
        new_file.write('\nRecurring')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.exts.r_sum_p))
        new_file.write('\nNegative\nOne-Time')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.exts.one_sum_n))
        new_file.write('\nRecurring')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.exts.r_sum_n))
        new_file.write('\nTotal: Present Expected Value\nBenefits')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.total_bens))
        new_file.write('\nCosts')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.total_costs))
        new_file.write('\nNet')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.net))
        any_ext = False
        for plan in self.plan_list:
            if len(plan.exts.indiv) > 0:
                any_ext = True
        if any_ext:
            new_file.write('\nNet with Externalities')
            for plan in self.plan_list:
                new_file.write(',$' + str(plan.net_w_ext))
        new_file.write('\nSavings-to-Investment Ratio')
        for plan in self.plan_list:
            new_file.write(',' + str(plan.sir()))
        new_file.write('\nInternal Rate of Return (%)')
        for plan in self.plan_list:
            new_file.write(',' + str(plan.irr()))
        new_file.write('\nReturn on Investment (%)')
        for plan in self.plan_list:
            new_file.write(',' + str(plan.roi()))
        new_file.write('\nNon-Disaster ROI (%)')
        for plan in self.plan_list:
            new_file.write(',' + str(plan.non_d_roi()))
        new_file.close()

    def csv_export_uncert(self):
        num_runs = 1000

        my_formats = [('Comma Separated Value', '*.csv'),]
        file_name = filedialog.asksaveasfilename(filetypes=my_formats, title="Save the file as...")
        if '.csv' != file_name[-4:]:
            file_name = file_name + '.csv'
        new_file = open(file_name, 'w')
        new_file.write('Outputs of Economic Evaluation: [' + self.title + ']\n')
        new_file.write('NOTE: All bounds on uncertainties are given with a ' + str(self.confidence)
                       + '% confidence interval. The number of runs was determined with a '
                       + str(self.tolerance) + '% tolerance.\n')
        for plan in self.plan_list:
            new_file.write('For ' + plan.name + ' (Alternative ' + str(plan.id_assign) + ') '
                           + str(plan.mc_iters) + ' Monte-Carlo simulations were run.\n')
        new_file.write('The random number seed for these runs was ' + str(self.seed) + '.\n')
        new_file.write(',Base Case,Lower Bound,Upper Bound')
        for i in range(1, self.num_plans):
            new_file.write(',Alternative ' + str(i) + ',Lower Bound,Upper Bound')
        new_file.write('\n,Base')
        for i in range(1, self.num_plans):
            new_file.write(',,,' + self.plan_list[i].name)
        new_file.write('\nBenefits\nDisaster Economic Benefits\nResponse and Recovery Costs')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.bens.r_sum))
            new_file.write(',$' + str(plan.bens.res_rec_range[0])
                           + ',$' + str(plan.bens.res_rec_range[1]))
        new_file.write('\nDirect Loss Reduction')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.bens.d_sum))
            new_file.write(',$' + str(plan.bens.direct_range[0])
                           + ',$' + str(plan.bens.direct_range[1]))
        new_file.write('\nIndirect Losses')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.bens.i_sum))
            new_file.write(',$' + str(plan.bens.indirect_range[0])
                           + ',$' + str(plan.bens.indirect_range[1]))
        new_file.write('\nDisaster Non-Market Benefits\nValue of Statistical Lives Saved')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.fat.stat_value_averted))
            new_file.write(',$' + str(plan.fat.value_range[0])
                           + ',$' + str(plan.fat.value_range[1]))
        new_file.write('\nNumber of Statistical Lives Saved')
        for plan in self.plan_list:
            new_file.write(',' + str(plan.fat.stat_averted))
            new_file.write(',' + str(plan.fat.num_range[0]) + ',' + str(plan.fat.num_range[1]))
        new_file.write('\nNon-disaster Related Benefits\nOne-Time')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.nond_bens.one_sum))
            new_file.write(',$' + str(plan.nond_bens.one_range[0])
                           + ',$' + str(plan.nond_bens.one_range[1]))
        new_file.write('\nRecurring')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.nond_bens.r_sum))
            new_file.write(',$' + str(plan.nond_bens.r_range[0])
                           + ',$' + str(plan.nond_bens.r_range[1]))
        new_file.write('\nCosts\nDirect Costs')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.costs.d_sum))
            new_file.write(',$' + str(plan.costs.direct_range[0])
                           + ',$' + str(plan.costs.direct_range[1]))
        new_file.write('\nIndirect Costs')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.costs.i_sum))
            new_file.write(',$' + str(plan.costs.indirect_range[0])
                           + ',$' + str(plan.costs.indirect_range[1]))
        new_file.write('\nOMR\nOne-Time')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.costs.omr_1_sum))
            new_file.write(',$' + str(plan.costs.omr_one_range[0])
                           + ',$' + str(plan.costs.omr_one_range[1]))
        new_file.write('\nRecurring')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.costs.omr_r_sum))
            new_file.write(',$' + str(plan.costs.omr_r_range[0])
                           + ',$' + str(plan.costs.omr_r_range[1]))
        new_file.write('\nExternalities\nPositive\nOne-Time')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.exts.one_sum_p))
            new_file.write(',$' + str(plan.exts.one_p_range[0])
                           + ',$' + str(plan.exts.one_p_range[1]))
        new_file.write('\nRecurring')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.exts.r_sum_p))
            new_file.write(',$' + str(plan.exts.r_p_range[0]) + ',$' + str(plan.exts.r_p_range[1]))
        new_file.write('\nNegative\nOne-Time')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.exts.one_sum_n))
            new_file.write(',$' + str(plan.exts.one_n_range[0])
                           + ',$' + str(plan.exts.one_n_range[1]))
        new_file.write('\nRecurring')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.exts.r_sum_n))
            new_file.write(',$' + str(plan.exts.r_n_range[0]) + ',$' + str(plan.exts.r_n_range[1]))
        new_file.write('\nTotal: Present Expected Value\nBenefits')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.total_bens))
            new_file.write(',$' + str(plan.ben_range[0]) + ',$' + str(plan.ben_range[1]))
        new_file.write('\nCosts')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.total_costs))
            new_file.write(',$' + str(plan.cost_range[0]) + ',$' + str(plan.cost_range[1]))
        new_file.write('\nNet')
        for plan in self.plan_list:
            new_file.write(',$' + str(plan.net))
            new_file.write(',$' + str(plan.net_range[0]) + ',$' + str(plan.net_range[1]))
        any_ext = False
        for plan in self.plan_list:
            if len(plan.exts.indiv) > 0:
                any_ext = True
        if any_ext:
            new_file.write('\nNet with Externalities')
            for plan in self.plan_list:
                new_file.write(',$' + str(plan.net_w_ext))
                new_file.write(',$' + str(plan.net_ext_range[0])
                               + ',$' + str(plan.net_ext_range[1]))
        new_file.write('\nSavings-to-Investment Ratio')
        for plan in self.plan_list:
            new_file.write(',' + str(plan.sir()))
            new_file.write(',' + str(plan.sir_range[0]) + ',' + str(plan.sir_range[1]))
        new_file.write('\nInternal Rate of Return (%)')
        for plan in self.plan_list:
            new_file.write(',' + str(plan.irr()))
            new_file.write(',' + str(plan.irr_range[0]) + ',' + str(plan.irr_range[1]))
        new_file.write('\nReturn on Investment (%)')
        for plan in self.plan_list:
            new_file.write(',' + str(plan.roi()))
            new_file.write(',' + str(plan.roi_range[0]) + ',' + str(plan.roi_range[1]))
        new_file.write('\nNon-Disaster ROI (%)')
        for plan in self.plan_list:
            new_file.write(',' + str(plan.non_d_roi()))
            new_file.write(',' + str(plan.nond_roi_range[0]) + ',' + str(plan.nond_roi_range[1]))
        new_file.close()

    def word_export(self):
        my_formats = [('Microsoft Word Document', '*.docx'),]
        file_name = filedialog.asksaveasfilename(filetypes=my_formats, title="Save the file as...")
        if '.docx' != file_name[-5:]:
            file_name = file_name + '.docx'

        any_ext = False
        for plan in self.plan_list:
            if len(plan.exts.indiv) > 0:
                any_ext = True

        doc = docx.Document()
        doc.add_heading('Economic Evaluation Complete Report\n' + self.title, 0)

        doc.add_heading('Analysis Base Information\n', 1)
        doc.add_paragraph('Number of Alternatives: ' + str(self.num_plans-1), style='ListBullet')
        doc.add_paragraph('Planning Horizon: ' + str(self.horizon) + ' years', style='ListBullet')
        doc.add_paragraph('Discount Rate: ' + str(self.discount_rate) + '%', style='ListBullet')
        doc.add_paragraph()
        doc.add_paragraph('Disaster Rate: Every ' + str(self.get_disaster_rate()[0]) + ' years',
                          style='ListBullet')
        doc.add_paragraph('Disaster Magnitude: ' + str(self.get_disaster_magnitude()[0])
                          + '% of build cost', style='ListBullet')
        doc.add_paragraph('Risk Preference: ' + str(self.risk_pref), style='ListBullet')
        doc.add_paragraph()
        doc.add_paragraph('Statistical Value of a Life: '+'${:.0f}'.format(float(self.stat_life)),
                          style='ListBullet')

        if any_ext:
            header_list = ['Plan Title', 'Total Benefits ($)', 'Total Costs ($)', 'Net ($)',
                           'Net with externalities ($)', 'SIR (%)', 'IRR (%)', 'ROI (%)',
                           'Non-Disaster ROI (%)']
        else:
            header_list = ['Plan Title', 'Total Benefits ($)', 'Total Costs ($)', 'Net ($)',
                           'SIR (%)', 'IRR (%)', 'ROI (%)', 'Non-Disaster ROI (%)']

        doc.add_heading('Summary\n', 1)
        sum_table = doc.add_table(rows = len(self.plan_list) + 1, cols=len(header_list),
                                  style='Light List Accent 1')

        for i in range(len(header_list)):
            sum_table.cell(0, i).text = header_list[i]
            sum_table.cell(0, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        sum_index = 1
        for plan in self.plan_list:
            sum_table.cell(sum_index, 0).text = plan.name
            sum_table.cell(sum_index, 1).text = '{:,.0f}'.format(plan.total_bens)
            sum_table.cell(sum_index, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sum_table.cell(sum_index, 2).text = '{:,.0f}'.format(plan.total_costs)
            sum_table.cell(sum_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sum_table.cell(sum_index, 3).text = '{:,.0f}'.format(plan.net)
            sum_table.cell(sum_index, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if any_ext:
                sum_table.cell(sum_index, 4).text = '{:,.0f}'.format(plan.net_w_ext)
                sum_table.cell(sum_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                sum_table.cell(sum_index, 5).text = '{:,.2f}'.format(plan.sir())
                sum_table.cell(sum_index, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if isinstance(plan.irr(), str):
                    sum_table.cell(sum_index, 6).text = plan.irr()
                    sum_table.cell(sum_index, 6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    sum_table.cell(sum_index, 6).text = '{:,.2f}'.format(plan.irr())
                    sum_table.cell(sum_index, 6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                sum_table.cell(sum_index, 7).text = '{:,.2f}'.format(plan.roi())
                sum_table.cell(sum_index, 7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                sum_table.cell(sum_index, 8).text = '{:,.2f}'.format(plan.non_d_roi())
                sum_table.cell(sum_index, 8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                sum_index += 1
            else:
                sum_table.cell(sum_index, 4).text = '{:,.2f}'.format(plan.sir())
                sum_table.cell(sum_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if isinstance(plan.irr(), str):
                    sum_table.cell(sum_index, 5).text = plan.irr()
                    sum_table.cell(sum_index, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    sum_table.cell(sum_index, 5).text = '{:,.2f}'.format(plan.irr())
                    sum_table.cell(sum_index, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                sum_table.cell(sum_index, 6).text = '{:,.2f}'.format(plan.roi())
                sum_table.cell(sum_index, 6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                sum_table.cell(sum_index, 7).text = '{:,.2f}'.format(plan.non_d_roi())
                sum_table.cell(sum_index, 7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                sum_index += 1

        for plan in self.plan_list:
            doc.add_heading(plan.name, 1)
            doc.add_heading('Alternative ' + str(plan.id_assign), 3)
            doc.add_heading('Fatalities Averted\n', 2)
            doc.add_paragraph('Number of Statistical Lives Saved: '
                              + '{:,.2f}'.format(plan.fat.stat_averted))
            doc.add_paragraph('Value of Statistical Lives Saved: '
                              + '${:,.0f}'.format(plan.fat.stat_value_averted))
            if plan.fat.desc != 'N/A':
                doc.add_paragraph('Description: ' + plan.fat.desc)
            doc.add_heading('Disaster-Related Benefits\n', 2)
            # == BENEFITS
            ben_table = doc.add_table(rows=len(plan.bens.indiv)+5, cols=3,
                                      style='Light List Accent 1')
            ben_index = 2
            ben_table.cell(0, 0).text = 'Title'
            ben_table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ben_table.cell(0, 1).text = 'Amount ($)'
            ben_table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ben_table.cell(0, 2).text = 'Effective Present Value ($)'
            ben_table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ben_table.cell(1, 0).text = 'Response and Recovery Cost Reductions'
            ben_table.cell(1, 2).text = '{:,.0f}'.format(plan.bens.r_sum)
            ben_table.cell(1, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for ben in plan.bens.indiv:
                if ben.ben_type == "res-rec":
                    ben_table.cell(ben_index, 0).text = ben.title
                    ben_table.cell(ben_index, 1).text = '{:,.0f}'.format(ben.amount)
                    ben_table.cell(ben_index, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ben_table.cell(ben_index, 2).text = '{:,.0f}'.format(plan.bens.on_dis_occ(ben.amount, plan.horizon, plan.recurrence, self.discount_rate))
                    ben_table.cell(ben_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ben_index += 1
            ben_table.cell(ben_index, 0).text = 'Direct Losses Prevented'
            ben_table.cell(ben_index, 2).text = '{:,.0f}'.format(plan.bens.d_sum)
            ben_table.cell(ben_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            ben_index += 1
            for ben in plan.bens.indiv:
                if ben.ben_type == "direct":
                    ben_table.cell(ben_index, 0).text = ben.title
                    ben_table.cell(ben_index, 1).text = '{:,.0f}'.format(ben.amount)
                    ben_table.cell(ben_index, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ben_table.cell(ben_index, 2).text = '{:,.0f}'.format(plan.bens.on_dis_occ(ben.amount, plan.horizon, plan.recurrence, self.discount_rate))
                    ben_table.cell(ben_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ben_index += 1
            ben_table.cell(ben_index, 0).text = 'Indirect Losses Prevented'
            ben_table.cell(ben_index, 2).text = '{:,.0f}'.format(plan.bens.i_sum)
            ben_table.cell(ben_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            ben_index += 1
            for ben in plan.bens.indiv:
                if ben.ben_type == "indirect":
                    ben_table.cell(ben_index, 0).text = ben.title
                    ben_table.cell(ben_index, 1).text = '{:,.0f}'.format(ben.amount)
                    ben_table.cell(ben_index, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ben_table.cell(ben_index, 2).text = '{:,.0f}'.format(plan.bens.on_dis_occ(ben.amount, plan.horizon, plan.recurrence, self.discount_rate))
                    ben_table.cell(ben_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ben_index += 1
            ben_table.cell(ben_index, 0).text = 'Total'
            ben_table.cell(ben_index, 2).text = '{:,.0f}'.format(plan.bens.total)
            ben_table.cell(ben_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            for ben in plan.bens.indiv:
                if ben.desc != 'N/A':
                    doc.add_paragraph(ben.title + ': ' + ben.desc)

            # == NON-D BENEFITS
            doc.add_heading('Resilience Dividend\n', 2)
            ben_table = doc.add_table(rows=len(plan.nond_bens.indiv)+4, cols=5,
                                      style='Light List Accent 1')
            ben_index = 2
            # Header
            ben_table.cell(0, 0).text = 'Title'
            ben_table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ben_table.cell(0, 1).text = 'Start Year'
            ben_table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ben_table.cell(0, 2).text = 'Recurrence (Years)'
            ben_table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ben_table.cell(0, 3).text = 'Amount ($)'
            ben_table.cell(0, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ben_table.cell(0, 4).text = 'Effective Present Value ($)'
            ben_table.cell(0, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ben_table.cell(1, 0).text = 'One Time Cost Reductions'
            ben_table.cell(1, 4).text = '{:,.0f}'.format(plan.nond_bens.one_sum)
            ben_table.cell(1, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for ben in plan.nond_bens.indiv:
                if ben.ben_type == "one-time":
                    ben_table.cell(ben_index, 0).text = ben.title
                    ben_table.cell(ben_index, 1).text = str(ben.times[0])
                    ben_table.cell(ben_index, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ben_table.cell(ben_index, 2).text = 'N/A'
                    ben_table.cell(ben_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ben_table.cell(ben_index, 3).text = '{:,.0f}'.format(float(ben.amount))
                    ben_table.cell(ben_index, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ben_table.cell(ben_index, 4).text = '{:,.0f}'.format(plan.nond_bens.calc_one_time(ben.amount, ben.times[0]))
                    ben_table.cell(ben_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ben_index += 1
            ben_table.cell(ben_index, 0).text = 'Recurring Cost Reductions'
            ben_table.cell(ben_index, 4).text = '{:,.0f}'.format(plan.nond_bens.r_sum)
            ben_table.cell(ben_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            ben_index += 1
            for ben in plan.nond_bens.indiv:
                if ben.ben_type == "recurring":
                    ben_table.cell(ben_index, 0).text = ben.title
                    ben_table.cell(ben_index, 1).text = str(ben.times[0])
                    ben_table.cell(ben_index, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ben_table.cell(ben_index, 2).text = str(ben.times[1])
                    ben_table.cell(ben_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ben_table.cell(ben_index, 3).text = '{:,.0f}'.format(float(ben.amount))
                    ben_table.cell(ben_index, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ben_table.cell(ben_index, 4).text = '{:,.0f}'.format(plan.nond_bens.calc_recur(ben.amount, ben.times[0], ben.times[1]))
                    ben_table.cell(ben_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ben_index += 1
            ben_table.cell(ben_index, 0).text = 'Total'
            ben_table.cell(ben_index, 4).text = '{:,.0f}'.format(plan.nond_bens.total)
            ben_table.cell(ben_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for ben in plan.nond_bens.indiv:
                if ben.desc != 'N/A':
                    doc.add_paragraph(ben.title + ': ' + ben.desc)

            # == COSTS
            doc.add_heading('Costs\n', 2)
            cost_table = doc.add_table(rows=len(plan.costs.indiv)+6, cols=5,
                                       style='Light List Accent 1')
            cost_index = 2
            # Header
            cost_table.cell(0, 0).text = 'Title'
            cost_table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cost_table.cell(0, 1).text = 'Start Year'
            cost_table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cost_table.cell(0, 2).text = 'Recurrence (Years)'
            cost_table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cost_table.cell(0, 3).text = 'Amount ($)'
            cost_table.cell(0, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cost_table.cell(0, 4).text = 'Effective Present Value ($)'
            cost_table.cell(0, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cost_table.cell(1, 0).text = 'Direct Costs'
            cost_table.cell(1, 4).text = '{:,.0f}'.format(plan.costs.d_sum)
            cost_table.cell(1, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for cost in plan.costs.indiv:
                if cost.cost_type == "direct":
                    cost_table.cell(cost_index, 0).text = cost.title
                    cost_table.cell(cost_index, 1).text = 'Start-Up'
                    cost_table.cell(cost_index, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    cost_table.cell(cost_index, 2).text = 'N/A'
                    cost_table.cell(cost_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    cost_table.cell(cost_index, 3).text = '{:,.0f}'.format(cost.amount)
                    cost_table.cell(cost_index, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    cost_table.cell(cost_index, 4).text = '{:,.0f}'.format(plan.costs.calc_one_time(cost.amount, cost.times[0]))
                    cost_table.cell(cost_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    cost_index += 1
            cost_table.cell(cost_index, 0).text = 'Indirect Costs'
            cost_table.cell(cost_index, 4).text = '{:,.0f}'.format(plan.costs.i_sum)
            cost_table.cell(cost_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cost_index += 1
            for cost in plan.costs.indiv:
                if cost.cost_type == "indirect":
                    cost_table.cell(cost_index, 0).text = cost.title
                    cost_table.cell(cost_index, 1).text = 'Start-Up'
                    cost_table.cell(cost_index, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    cost_table.cell(cost_index, 2).text = 'N/A'
                    cost_table.cell(cost_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    cost_table.cell(cost_index, 3).text = '{:,.0f}'.format(cost.amount)
                    cost_table.cell(cost_index, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    cost_table.cell(cost_index, 4).text = '{:,.0f}'.format(plan.costs.calc_one_time(cost.amount, cost.times[0]))
                    cost_table.cell(cost_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    cost_index += 1
            cost_table.cell(cost_index, 0).text = 'OMR Costs: One-Time'
            cost_table.cell(cost_index, 4).text = '{:,.0f}'.format(plan.costs.omr_1_sum)
            cost_table.cell(cost_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cost_index += 1
            for cost in plan.costs.indiv:
                if (cost.cost_type == "omr") & (cost.omr_type == "one-time"):
                    cost_table.cell(cost_index, 0).text = cost.title
                    cost_table.cell(cost_index, 1).text = str(cost.times[0])
                    cost_table.cell(cost_index, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    cost_table.cell(cost_index, 2).text = 'N/A'
                    cost_table.cell(cost_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    cost_table.cell(cost_index, 3).text = '{:,.0f}'.format(cost.amount)
                    cost_table.cell(cost_index, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    cost_table.cell(cost_index, 4).text = '{:,.0f}'.format(plan.costs.calc_one_time(cost.amount, float(cost.times[0])))
                    cost_table.cell(cost_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    cost_index += 1
            cost_table.cell(cost_index, 0).text = 'OMR Costs: Recurring'
            cost_table.cell(cost_index, 4).text = '{:,.0f}'.format(plan.costs.omr_r_sum)
            cost_table.cell(cost_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cost_index += 1
            for cost in plan.costs.indiv:
                if (cost.cost_type == "omr") & (cost.omr_type == "recurring"):
                    cost_table.cell(cost_index, 0).text = cost.title
                    cost_table.cell(cost_index, 1).text = str(cost.times[0])
                    cost_table.cell(cost_index, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    cost_table.cell(cost_index, 2).text = str(cost.times[1])
                    cost_table.cell(cost_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    cost_table.cell(cost_index, 3).text = '{:,.0f}'.format(cost.amount)
                    cost_table.cell(cost_index, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    cost_table.cell(cost_index, 4).text = '{:,.0f}'.format(plan.costs.calc_recur(cost.amount, cost.times[0], cost.times[1]))
                    cost_table.cell(cost_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    cost_index += 1
            cost_table.cell(cost_index, 0).text = 'Total'
            cost_table.cell(cost_index, 4).text = '{:,.0f}'.format(plan.costs.total)
            cost_table.cell(cost_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for cost in plan.costs.indiv:
                if cost.desc != 'N/A':
                    doc.add_paragraph(cost.title + ': ' + cost.desc)

            # == EXTERNALITIES
            doc.add_heading('Externalities\n', 2)
            ext_table = doc.add_table(rows=len(plan.exts.indiv)+6, cols=5, style='Light List Accent 1')
            #ext_table.style = 'TableGrid'
            ext_index = 2
            # Header
            ext_table.cell(0, 0).text = 'Title'
            ext_table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ext_table.cell(0, 1).text = 'Start Year'
            ext_table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ext_table.cell(0, 2).text = 'Recurrence (Years)'
            ext_table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ext_table.cell(0, 3).text = 'Amount ($)'
            ext_table.cell(0, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ext_table.cell(0, 4).text = 'Effective Present Value ($)'
            ext_table.cell(0, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ext_table.cell(1, 0).text = 'One Time Positive Externalities'
            ext_table.cell(1, 4).text = '{:,.0f}'.format(plan.exts.one_sum_p)
            ext_table.cell(1, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for ext in plan.exts.indiv:
                if (ext.ext_type == "one-time") & (ext.pm == "+"):
                    ext_table.cell(ext_index, 0).text = ext.title
                    ext_table.cell(ext_index, 1).text = str(ext.times[0])
                    ext_table.cell(ext_index, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ext_table.cell(ext_index, 2).text = 'N/A'
                    ext_table.cell(ext_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ext_table.cell(ext_index, 3).text = '{:,.0f}'.format(float(ext.amount))
                    ext_table.cell(ext_index, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ext_table.cell(ext_index, 4).text = '{:,.0f}'.format(plan.exts.calc_one_time(ext.amount, ext.times[0]))
                    ext_table.cell(ext_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ext_index += 1
            ext_table.cell(ext_index, 0).text = 'Recurring Positive Externalities'
            ext_table.cell(ext_index, 4).text = '{:,.0f}'.format(plan.exts.r_sum_p)
            ext_table.cell(ext_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            ext_index += 1
            for ext in plan.exts.indiv:
                if (ext.ext_type == "recurring") & (ext.pm == "+"):
                    ext_table.cell(ext_index, 0).text = ext.title
                    ext_table.cell(ext_index, 1).text = str(ext.times[0])
                    ext_table.cell(ext_index, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ext_table.cell(ext_index, 2).text = str(ext.times[1])
                    ext_table.cell(ext_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ext_table.cell(ext_index, 3).text = '{:,.0f}'.format(float(ext.amount))
                    ext_table.cell(ext_index, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ext_table.cell(ext_index, 4).text = '{:,.0f}'.format(plan.exts.calc_recur(ext.amount, ext.times[0], ext.times[1]))
                    ext_table.cell(ext_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ext_index += 1
            ext_table.cell(ext_index, 0).text = 'One Time Negative Externalities'
            ext_table.cell(ext_index, 4).text = '{:,.0f}'.format(plan.exts.one_sum_n)
            ext_table.cell(ext_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            ext_index += 1
            for ext in plan.exts.indiv:
                if (ext.ext_type == "one-time") & (ext.pm == "-"):
                    ext_table.cell(ext_index, 0).text = ext.title
                    ext_table.cell(ext_index, 1).text = str(ext.times[0])
                    ext_table.cell(ext_index, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ext_table.cell(ext_index, 2).text = 'N/A'
                    ext_table.cell(ext_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ext_table.cell(ext_index, 3).text = '{:,.0f}'.format(float(ext.amount))
                    ext_table.cell(ext_index, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ext_table.cell(ext_index, 4).text = '{:,.0f}'.format(plan.exts.calc_one_time(ext.amount, ext.times[0]))
                    ext_table.cell(ext_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ext_index += 1
            ext_table.cell(ext_index, 0).text = 'Recurring Negative Externalities'
            ext_table.cell(ext_index, 4).text = '{:,.0f}'.format(plan.exts.r_sum_n)
            ext_table.cell(ext_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            ext_index += 1
            for ext in plan.exts.indiv:
                if (ext.ext_type == "recurring") & (ext.pm == "-"):
                    ext_table.cell(ext_index, 0).text = ext.title
                    ext_table.cell(ext_index, 1).text = str(ext.times[0])
                    ext_table.cell(ext_index, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ext_table.cell(ext_index, 2).text = str(ext.times[1])
                    ext_table.cell(ext_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ext_table.cell(ext_index, 3).text = '{:,.0f}'.format(float(ext.amount))
                    ext_table.cell(ext_index, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ext_table.cell(ext_index, 4).text = '{:,.0f}'.format(plan.exts.calc_recur(ext.amount, ext.times[0], ext.times[1]))
                    ext_table.cell(ext_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ext_index += 1
            ext_table.cell(ext_index, 0).text = 'Total'
            ext_table.cell(ext_index, 4).text = '{:,.0f}'.format(plan.exts.total_p - plan.exts.total_n)
            ext_table.cell(ext_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for ext in plan.exts.indiv:
                if ext.desc != 'N/A':
                    doc.add_paragraph(ext.title + ': ' + ext.desc)

        sections = doc.sections
        for section in sections:
            section.left_margin = 914400
            section.right_margin = 914400
        doc.save(file_name)

    def uncert_string(self, uncert, values):
        """Creates a nice string given a certain uncertainty type and the associated values."""
        values = list(values)
        try:
            if (uncert in {'tri', 'rect'}) & (str(values[2]) == str(0)):
                values[2] = values[1]
        except IndexError:
            pass
        if uncert == "none":
            return "N/A"
        elif uncert == "gauss":
            return "Gaussian distribution with variance of " + str(values[0])
        elif uncert == "tri":
            return "Triangular distribution with a min of " + str(values[0]) + " and a max of " + str(values[2])
        elif uncert == "rect":
            return "Rectangular distribution with a min of " + str(values[0]) + " and a max of " + str(values[2])
        elif uncert == "discrete":
            return str(values[3]) + "% chance of " + str(values[0]) + ", " + str(values[4]) + "% chance of " + str(values[1]) + ", and " + str(values[5]) + "% chance of " + str(values[2])

    def word_export_uncert(self):
        my_formats = [('Microsoft Word Document', '*.docx'),]
        file_name = filedialog.asksaveasfilename(filetypes=my_formats, title="Save the file as...")
        if '.docx' != file_name[-5:]:
            file_name = file_name + '.docx'

        any_ext = False
        for plan in self.plan_list:
            if len(plan.exts.indiv) > 0:
                any_ext = True

        doc = docx.Document()
        doc.add_heading('Economic Evaluation Complete Report\n' + self.title, 0)
        doc.add_paragraph('NOTE: All bounds on uncertainties are given with a ' + str(self.confidence)
                        + '% confidence interval. The number of runs was determined with a '
                        + str(self.tolerance) + '% tolerance.')
        for plan in self.plan_list:
            doc.add_paragraph('For ' + plan.name + ' (Alternative ' + str(plan.id_assign) + ') '
                           + str(plan.mc_iters) + ' Monte-Carlo simulations were run.')
        doc.add_paragraph('The random number seed for these runs was ' + str(self.seed) + '.')


        doc.add_heading('Analysis Base Information\n', 1)
        doc.add_paragraph('Number of Alternatives: ' + str(self.num_plans-1), style='ListBullet')
        doc.add_paragraph('Planning Horizon: ' + str(self.horizon) + ' years', style='ListBullet')
        doc.add_paragraph('Discount Rate: ' + str(self.discount_rate) + '%', style='ListBullet')
        doc.add_paragraph()
        doc.add_paragraph('Disaster Rate: Every ' + str(self.get_disaster_rate()[0]) + ' years', style='ListBullet')
        doc.add_paragraph('Uncertainty in Disaster Rate: ' + self.uncert_string(self.get_disaster_rate()[2], self.get_disaster_rate()[1]), style='ListBullet')
        doc.add_paragraph('Disaster Magnitude: ' + str(self.get_disaster_magnitude()[0]) + '% of build cost', style='ListBullet')
        doc.add_paragraph('Uncertainty in Disaster Magnitude: ' + self.uncert_string(self.get_disaster_magnitude()[2], self.get_disaster_magnitude()[1]), style='ListBullet')
        doc.add_paragraph('Risk Preference: ' + str(self.risk_pref), style='ListBullet')
        doc.add_paragraph()
        doc.add_paragraph('Statistical Value of a Life: '+'${:.0f}'.format(float(self.stat_life)), style='ListBullet')

        if any_ext:
            header_list = ['Plan Title', 'Total Benefits ($)', 'Total Costs ($)', 'Net ($)', 'Net with externalities ($)', 'SIR (%)', 'IRR (%)', 'ROI (%)', 'Non-Disaster ROI (%)']
        else:
            header_list = ['Plan Title', 'Total Benefits ($)', 'Total Costs ($)', 'Net ($)', 'SIR (%)', 'IRR (%)', 'ROI (%)', 'Non-Disaster ROI (%)']

        doc.add_heading('Summary\n', 1)
        sum_table = doc.add_table(rows = len(self.plan_list) + 1, cols=len(header_list), style='Light List Accent 1')

        for i in range(len(header_list)):
            sum_table.cell(0, i).text = header_list[i]
            sum_table.cell(0, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        sum_index = 1
        for plan in self.plan_list:
            sum_table.cell(sum_index, 0).text = plan.name
            sum_table.cell(sum_index, 1).text = '{:,.0f}'.format(plan.total_bens)
            sum_table.cell(sum_index, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sum_table.cell(sum_index, 2).text = '{:,.0f}'.format(plan.total_costs)
            sum_table.cell(sum_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sum_table.cell(sum_index, 3).text = '{:,.0f}'.format(plan.net)
            sum_table.cell(sum_index, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if any_ext:
                sum_table.cell(sum_index, 4).text = '{:,.0f}'.format(plan.net_w_ext)
                sum_table.cell(sum_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                sum_table.cell(sum_index, 5).text = '{:,.2f}'.format(plan.sir())
                sum_table.cell(sum_index, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if isinstance(plan.irr(), str):
                    sum_table.cell(sum_index, 6).text = plan.irr()
                    sum_table.cell(sum_index, 6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    sum_table.cell(sum_index, 6).text = '{:,.2f}'.format(plan.irr())
                    sum_table.cell(sum_index, 6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                sum_table.cell(sum_index, 7).text = '{:,.2f}'.format(plan.roi())
                sum_table.cell(sum_index, 7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                sum_table.cell(sum_index, 8).text = '{:,.2f}'.format(plan.non_d_roi())
                sum_table.cell(sum_index, 8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                sum_index += 1
            else:
                sum_table.cell(sum_index, 4).text = '{:,.2f}'.format(plan.sir())
                sum_table.cell(sum_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if isinstance(plan.irr(), str):
                    sum_table.cell(sum_index, 5).text = plan.irr()
                    sum_table.cell(sum_index, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    sum_table.cell(sum_index, 5).text = '{:,.2f}'.format(plan.irr())
                    sum_table.cell(sum_index, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                sum_table.cell(sum_index, 6).text = '{:,.2f}'.format(plan.roi())
                sum_table.cell(sum_index, 6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                sum_table.cell(sum_index, 7).text = '{:,.2f}'.format(plan.non_d_roi())
                sum_table.cell(sum_index, 7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                sum_index += 1

        for plan in self.plan_list:
            doc.add_heading(plan.name, 1)
            doc.add_heading('Alternative ' + str(plan.id_assign), 3)
            doc.add_heading('Fatalities Averted\n', 2)
            doc.add_paragraph('Number of Statistical Lives Saved: ' + '{:,.2f}'.format(plan.fat.stat_averted))
            doc.add_paragraph('Value of Statistical Lives Saved: ' + '${:,.0f}'.format(plan.fat.stat_value_averted))
            if plan.fat.desc != 'N/A':
                doc.add_paragraph('Description: ' + plan.fat.desc)
            doc.add_heading('Disaster-Related Benefits\n', 2)
            # == BENEFITS
            ben_table = doc.add_table(rows=len(plan.bens.indiv)+5, cols=3, style='Light List Accent 1')
            ben_index = 2
            ben_table.cell(0, 0).text = 'Title'
            ben_table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ben_table.cell(0, 1).text = 'Amount ($)'
            ben_table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ben_table.cell(0, 2).text = 'Effective Present Value ($)'
            ben_table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ben_table.cell(1, 0).text = 'Response and Recovery Cost Reductions'
            ben_table.cell(1, 2).text = '{:,.0f}'.format(plan.bens.r_sum)
            ben_table.cell(1, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for ben in plan.bens.indiv:
                if ben.ben_type == "res-rec":
                    ben_table.cell(ben_index, 0).text = ben.title
                    ben_table.cell(ben_index, 1).text = '{:,.0f}'.format(ben.amount)
                    ben_table.cell(ben_index, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ben_table.cell(ben_index, 2).text = '{:,.0f}'.format(plan.bens.on_dis_occ(ben.amount, plan.horizon, plan.recurrence, self.discount_rate))
                    ben_table.cell(ben_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ben_index += 1
            ben_table.cell(ben_index, 0).text = 'Direct Losses Prevented'
            ben_table.cell(ben_index, 2).text = '{:,.0f}'.format(plan.bens.d_sum)
            ben_table.cell(ben_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            ben_index += 1
            for ben in plan.bens.indiv:
                if ben.ben_type == "direct":
                    ben_table.cell(ben_index, 0).text = ben.title
                    ben_table.cell(ben_index, 1).text = '{:,.0f}'.format(ben.amount)
                    ben_table.cell(ben_index, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ben_table.cell(ben_index, 2).text = '{:,.0f}'.format(plan.bens.on_dis_occ(ben.amount, plan.horizon, plan.recurrence, self.discount_rate))
                    ben_table.cell(ben_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ben_index += 1
            ben_table.cell(ben_index, 0).text = 'Indirect Losses Prevented'
            ben_table.cell(ben_index, 2).text = '{:,.0f}'.format(plan.bens.i_sum)
            ben_table.cell(ben_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            ben_index += 1
            for ben in plan.bens.indiv:
                if ben.ben_type == "indirect":
                    ben_table.cell(ben_index, 0).text = ben.title
                    ben_table.cell(ben_index, 1).text = '{:,.0f}'.format(ben.amount)
                    ben_table.cell(ben_index, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ben_table.cell(ben_index, 2).text = '{:,.0f}'.format(plan.bens.on_dis_occ(ben.amount, plan.horizon, plan.recurrence, self.discount_rate))
                    ben_table.cell(ben_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ben_index += 1
            ben_table.cell(ben_index, 0).text = 'Total'
            ben_table.cell(ben_index, 2).text = '{:,.0f}'.format(plan.bens.total)
            ben_table.cell(ben_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            for ben in plan.bens.indiv:
                if ben.dist != 'none':
                    doc.add_paragraph(ben.title + ': ' + self.uncert_string(ben.dist, ben.range))
                if ben.desc != 'N/A':
                    doc.add_paragraph(ben.title + ': ' + ben.desc)

            # == NON-D BENEFITS
            doc.add_heading('Resilience Dividend\n', 2)
            ben_table = doc.add_table(rows=len(plan.nond_bens.indiv)+4, cols=5, style='Light List Accent 1')
            ben_index = 2
            # Header
            ben_table.cell(0, 0).text = 'Title'
            ben_table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ben_table.cell(0, 1).text = 'Start Year'
            ben_table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ben_table.cell(0, 2).text = 'Recurrence (Years)'
            ben_table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ben_table.cell(0, 3).text = 'Amount ($)'
            ben_table.cell(0, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ben_table.cell(0, 4).text = 'Effective Present Value ($)'
            ben_table.cell(0, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ben_table.cell(1, 0).text = 'One Time Cost Reductions'
            ben_table.cell(1, 4).text = '{:,.0f}'.format(plan.nond_bens.one_sum)
            ben_table.cell(1, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for ben in plan.nond_bens.indiv:
                if ben.ben_type == "one-time":
                    ben_table.cell(ben_index, 0).text = ben.title
                    ben_table.cell(ben_index, 1).text = str(ben.times[0])
                    ben_table.cell(ben_index, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ben_table.cell(ben_index, 2).text = 'N/A'
                    ben_table.cell(ben_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ben_table.cell(ben_index, 3).text = '{:,.0f}'.format(float(ben.amount))
                    ben_table.cell(ben_index, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ben_table.cell(ben_index, 4).text = '{:,.0f}'.format(plan.nond_bens.calc_one_time(ben.amount, ben.times[0]))
                    ben_table.cell(ben_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ben_index += 1
            ben_table.cell(ben_index, 0).text = 'Recurring Cost Reductions'
            ben_table.cell(ben_index, 4).text = '{:,.0f}'.format(plan.nond_bens.r_sum)
            ben_table.cell(ben_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            ben_index += 1
            for ben in plan.nond_bens.indiv:
                if ben.ben_type == "recurring":
                    ben_table.cell(ben_index, 0).text = ben.title
                    ben_table.cell(ben_index, 1).text = str(ben.times[0])
                    ben_table.cell(ben_index, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ben_table.cell(ben_index, 2).text = str(ben.times[1])
                    ben_table.cell(ben_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ben_table.cell(ben_index, 3).text = '{:,.0f}'.format(float(ben.amount))
                    ben_table.cell(ben_index, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ben_table.cell(ben_index, 4).text = '{:,.0f}'.format(plan.nond_bens.calc_recur(ben.amount, ben.times[0], ben.times[1]))
                    ben_table.cell(ben_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ben_index += 1
            ben_table.cell(ben_index, 0).text = 'Total'
            ben_table.cell(ben_index, 4).text = '{:,.0f}'.format(plan.nond_bens.total)
            ben_table.cell(ben_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for ben in plan.nond_bens.indiv:
                if ben.dist != 'none':
                    doc.add_paragraph(ben.title + ': ' + self.uncert_string(ben.dist, ben.range))
                if ben.desc != 'N/A':
                    doc.add_paragraph(ben.title + ': ' + ben.desc)

            # == COSTS
            doc.add_heading('Costs\n', 2)
            cost_table = doc.add_table(rows=len(plan.costs.indiv)+6, cols=5, style='Light List Accent 1')
            cost_index = 2
            # Header
            cost_table.cell(0, 0).text = 'Title'
            cost_table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cost_table.cell(0, 1).text = 'Start Year'
            cost_table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cost_table.cell(0, 2).text = 'Recurrence (Years)'
            cost_table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cost_table.cell(0, 3).text = 'Amount ($)'
            cost_table.cell(0, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cost_table.cell(0, 4).text = 'Effective Present Value ($)'
            cost_table.cell(0, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cost_table.cell(1, 0).text = 'Direct Costs'
            cost_table.cell(1, 4).text = '{:,.0f}'.format(plan.costs.d_sum)
            cost_table.cell(1, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for cost in plan.costs.indiv:
                if cost.cost_type == "direct":
                    cost_table.cell(cost_index, 0).text = cost.title
                    cost_table.cell(cost_index, 1).text = 'Start-Up'
                    cost_table.cell(cost_index, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    cost_table.cell(cost_index, 2).text = 'N/A'
                    cost_table.cell(cost_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    cost_table.cell(cost_index, 3).text = '{:,.0f}'.format(cost.amount)
                    cost_table.cell(cost_index, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    cost_table.cell(cost_index, 4).text = '{:,.0f}'.format(plan.costs.calc_one_time(cost.amount, cost.times[0]))
                    cost_table.cell(cost_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    cost_index += 1
            cost_table.cell(cost_index, 0).text = 'Indirect Costs'
            cost_table.cell(cost_index, 4).text = '{:,.0f}'.format(plan.costs.i_sum)
            cost_table.cell(cost_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cost_index += 1
            for cost in plan.costs.indiv:
                if cost.cost_type == "indirect":
                    cost_table.cell(cost_index, 0).text = cost.title
                    cost_table.cell(cost_index, 1).text = 'Start-Up'
                    cost_table.cell(cost_index, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    cost_table.cell(cost_index, 2).text = 'N/A'
                    cost_table.cell(cost_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    cost_table.cell(cost_index, 3).text = '{:,.0f}'.format(cost.amount)
                    cost_table.cell(cost_index, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    cost_table.cell(cost_index, 4).text = '{:,.0f}'.format(plan.costs.calc_one_time(cost.amount, cost.times[0]))
                    cost_table.cell(cost_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    cost_index += 1
            cost_table.cell(cost_index, 0).text = 'OMR Costs: One-Time'
            cost_table.cell(cost_index, 4).text = '{:,.0f}'.format(plan.costs.omr_1_sum)
            cost_table.cell(cost_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cost_index += 1
            for cost in plan.costs.indiv:
                if (cost.cost_type == "omr") & (cost.omr_type == "one-time"):
                    cost_table.cell(cost_index, 0).text = cost.title
                    cost_table.cell(cost_index, 1).text = str(cost.times[0])
                    cost_table.cell(cost_index, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    cost_table.cell(cost_index, 2).text = 'N/A'
                    cost_table.cell(cost_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    cost_table.cell(cost_index, 3).text = '{:,.0f}'.format(cost.amount)
                    cost_table.cell(cost_index, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    cost_table.cell(cost_index, 4).text = '{:,.0f}'.format(plan.costs.calc_one_time(cost.amount, float(cost.times[0])))
                    cost_table.cell(cost_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    cost_index += 1
            cost_table.cell(cost_index, 0).text = 'OMR Costs: Recurring'
            cost_table.cell(cost_index, 4).text = '{:,.0f}'.format(plan.costs.omr_r_sum)
            cost_table.cell(cost_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cost_index += 1
            for cost in plan.costs.indiv:
                if (cost.cost_type == "omr") & (cost.omr_type == "recurring"):
                    cost_table.cell(cost_index, 0).text = cost.title
                    cost_table.cell(cost_index, 1).text = str(cost.times[0])
                    cost_table.cell(cost_index, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    cost_table.cell(cost_index, 2).text = str(cost.times[1])
                    cost_table.cell(cost_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    cost_table.cell(cost_index, 3).text = '{:,.0f}'.format(cost.amount)
                    cost_table.cell(cost_index, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    cost_table.cell(cost_index, 4).text = '{:,.0f}'.format(plan.costs.calc_recur(cost.amount, cost.times[0], cost.times[1]))
                    cost_table.cell(cost_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    cost_index += 1
            cost_table.cell(cost_index, 0).text = 'Total'
            cost_table.cell(cost_index, 4).text = '{:,.0f}'.format(plan.costs.total)
            cost_table.cell(cost_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for cost in plan.costs.indiv:
                if cost.dist != 'none':
                    doc.add_paragraph(cost.title + ': ' + self.uncert_string(cost.dist, cost.range))
                if cost.desc != 'N/A':
                    doc.add_paragraph(cost.title + ': ' + cost.desc)

            # == EXTERNALITIES
            doc.add_heading('Externalities\n', 2)
            ext_table = doc.add_table(rows=len(plan.exts.indiv)+6, cols=5, style='Light List Accent 1')
            #ext_table.style = 'TableGrid'
            ext_index = 2
            # Header
            ext_table.cell(0, 0).text = 'Title'
            ext_table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ext_table.cell(0, 1).text = 'Start Year'
            ext_table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ext_table.cell(0, 2).text = 'Recurrence (Years)'
            ext_table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ext_table.cell(0, 3).text = 'Amount ($)'
            ext_table.cell(0, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ext_table.cell(0, 4).text = 'Effective Present Value ($)'
            ext_table.cell(0, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ext_table.cell(1, 0).text = 'One Time Positive Externalities'
            ext_table.cell(1, 4).text = '{:,.0f}'.format(plan.exts.one_sum_p)
            ext_table.cell(1, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for ext in plan.exts.indiv:
                if (ext.ext_type == "one-time") & (ext.pm == "+"):
                    ext_table.cell(ext_index, 0).text = ext.title
                    ext_table.cell(ext_index, 1).text = str(ext.times[0])
                    ext_table.cell(ext_index, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ext_table.cell(ext_index, 2).text = 'N/A'
                    ext_table.cell(ext_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ext_table.cell(ext_index, 3).text = '{:,.0f}'.format(float(ext.amount))
                    ext_table.cell(ext_index, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ext_table.cell(ext_index, 4).text = '{:,.0f}'.format(plan.exts.calc_one_time(ext.amount, ext.times[0]))
                    ext_table.cell(ext_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ext_index += 1
            ext_table.cell(ext_index, 0).text = 'Recurring Positive Externalities'
            ext_table.cell(ext_index, 4).text = '{:,.0f}'.format(plan.exts.r_sum_p)
            ext_table.cell(ext_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            ext_index += 1
            for ext in plan.exts.indiv:
                if (ext.ext_type == "recurring") & (ext.pm == "+"):
                    ext_table.cell(ext_index, 0).text = ext.title
                    ext_table.cell(ext_index, 1).text = str(ext.times[0])
                    ext_table.cell(ext_index, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ext_table.cell(ext_index, 2).text = str(ext.times[1])
                    ext_table.cell(ext_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ext_table.cell(ext_index, 3).text = '{:,.0f}'.format(float(ext.amount))
                    ext_table.cell(ext_index, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ext_table.cell(ext_index, 4).text = '{:,.0f}'.format(plan.exts.calc_recur(ext.amount, ext.times[0], ext.times[1]))
                    ext_table.cell(ext_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ext_index += 1
            ext_table.cell(ext_index, 0).text = 'One Time Negative Externalities'
            ext_table.cell(ext_index, 4).text = '{:,.0f}'.format(plan.exts.one_sum_n)
            ext_table.cell(ext_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            ext_index += 1
            for ext in plan.exts.indiv:
                if (ext.ext_type == "one-time") & (ext.pm == "-"):
                    ext_table.cell(ext_index, 0).text = ext.title
                    ext_table.cell(ext_index, 1).text = str(ext.times[0])
                    ext_table.cell(ext_index, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ext_table.cell(ext_index, 2).text = 'N/A'
                    ext_table.cell(ext_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ext_table.cell(ext_index, 3).text = '{:,.0f}'.format(float(ext.amount))
                    ext_table.cell(ext_index, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ext_table.cell(ext_index, 4).text = '{:,.0f}'.format(plan.exts.calc_one_time(ext.amount, ext.times[0]))
                    ext_table.cell(ext_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ext_index += 1
            ext_table.cell(ext_index, 0).text = 'Recurring Negative Externalities'
            ext_table.cell(ext_index, 4).text = '{:,.0f}'.format(plan.exts.r_sum_n)
            ext_table.cell(ext_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            ext_index += 1
            for ext in plan.exts.indiv:
                if (ext.ext_type == "recurring") & (ext.pm == "-"):
                    ext_table.cell(ext_index, 0).text = ext.title
                    ext_table.cell(ext_index, 1).text = str(ext.times[0])
                    ext_table.cell(ext_index, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ext_table.cell(ext_index, 2).text = str(ext.times[1])
                    ext_table.cell(ext_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ext_table.cell(ext_index, 3).text = '{:,.0f}'.format(float(ext.amount))
                    ext_table.cell(ext_index, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ext_table.cell(ext_index, 4).text = '{:,.0f}'.format(plan.exts.calc_recur(ext.amount, ext.times[0], ext.times[1]))
                    ext_table.cell(ext_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    ext_index += 1
            ext_table.cell(ext_index, 0).text = 'Total'
            ext_table.cell(ext_index, 4).text = '{:,.0f}'.format(plan.exts.total_p - plan.exts.total_n)
            ext_table.cell(ext_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for ext in plan.exts.indiv:
                if ext.dist != 'none':
                    doc.add_paragraph(ext.title + ': ' + self.uncert_string(ext.dist, ext.range))
                if ext.desc != 'N/A':
                    doc.add_paragraph(ext.title + ': ' + ext.desc)

        sections = doc.sections
        for section in sections:
            section.left_margin = 914400
            section.right_margin = 914400
        doc.save(file_name)


    def get_disaster_rate(self):
        """ Returns the disaster rate for the first plan
        since at the moment all plans will be the same."""
        return [self.plan_list[0].recurrence,
                self.plan_list[0].recurr_uncert,
                self.plan_list[0].recurr_dist]

    def get_disaster_magnitude(self):
        """ Returns the disaster magnitude for the first plan
        since at the moment all plans will be the same."""
        return [self.plan_list[0].magnitude,
                self.plan_list[0].mag_uncert,
                self.plan_list[0].mag_dist]

    def summer(self):
        """ Sums up all of the pieces."""
        for plan in self.plan_list:
            plan.sum_it(self.horizon)

    def monte(self, new_seed, confidence, tol, low_iters = 100, high_iters = 102400):
        """ Runs the monte-carlo everything."""
        self.seed = new_seed
        self.confidence = confidence
        self.tolerance = tol
        tol_percent = tol/100
        num_iters = low_iters
        ### NOTE: It's mad about this call, claiming it will pull an error. It doesn't
        np.random.seed(seed=new_seed)

        for plan in self.plan_list:
            ben_direct_totals = []
            ben_indirect_totals = []
            res_rec_totals = []
            fat_num_totals = []
            fat_value_totals = []
            cost_direct_totals = []
            cost_indirect_totals = []
            cost_omr_1_totals = []
            cost_omr_r_totals = []
            nond_1_totals = []
            nond_r_totals = []
            ext_p_one_totals = []
            ext_p_r_totals = []
            ext_n_one_totals = []
            ext_n_r_totals = []

            ben_totals = []
            cost_totals = []
            net_totals = []
            net_ext_totals = []
            irr_totals = []
            sir_totals = []
            roi_totals = []
            nond_roi_totals = []

            similar_list = []
            cost_ben_net = [False, False, False, False]
            old_iters = 0
            num_iters = low_iters
            while (not all(cost_ben_net)) and num_iters < high_iters:

                old_ben = list(plan.ben_range)
                old_cost = list(plan.cost_range)
                old_net = list(plan.net_range)
                old_ext_net = list(plan.net_ext_range)

                for i in range(old_iters, num_iters):
                    similar_list.append(self.one_iter(plan))
                    similar_list[i].sum_it(self.horizon)
                    ben_direct_totals.append(similar_list[i].bens.d_sum)
                    ben_indirect_totals.append(similar_list[i].bens.i_sum)
                    fat_num_totals.append(similar_list[i].fat.stat_averted)
                    fat_value_totals.append(similar_list[i].fat.stat_value_averted)
                    cost_direct_totals.append(similar_list[i].costs.d_sum)
                    cost_indirect_totals.append(similar_list[i].costs.i_sum)
                    cost_omr_1_totals.append(similar_list[i].costs.omr_1_sum)
                    cost_omr_r_totals.append(similar_list[i].costs.omr_r_sum)
                    res_rec_totals.append(similar_list[i].bens.r_sum)
                    nond_1_totals.append(similar_list[i].nond_bens.one_sum)
                    nond_r_totals.append(similar_list[i].nond_bens.r_sum)
                    ext_p_one_totals.append(similar_list[i].exts.one_sum_p)
                    ext_p_r_totals.append(similar_list[i].exts.r_sum_p)
                    ext_n_one_totals.append(similar_list[i].exts.one_sum_n)
                    ext_n_r_totals.append(similar_list[i].exts.r_sum_n)
                    ben_totals.append(similar_list[i].total_bens)
                    cost_totals.append(similar_list[i].total_costs)
                    net_totals.append(similar_list[i].net)
                    net_ext_totals.append(similar_list[i].net_w_ext)
                    if isinstance(similar_list[i].irr(), str):
                        irr_totals.append(-1)
                    else:
                        irr_totals.append(similar_list[i].irr())
                    sir_totals.append(similar_list[i].sir())
                    roi_totals.append(similar_list[i].roi())
                    nond_roi_totals.append(similar_list[i].non_d_roi())

                ben_totals.sort()
                cost_totals.sort()
                net_totals.sort()
                net_ext_totals.sort()

                first_num = math.floor(num_iters*(1-confidence/100)/2)
                last_num = math.ceil(num_iters - num_iters*(1-confidence/100)/2)#num_iters - first_num
                #print(first_num, similar_list.index(plan), last_num, num_iters)

                plan.ben_range = [ben_totals[first_num], ben_totals[last_num]]
                plan.cost_range = [cost_totals[first_num], cost_totals[last_num]]
                plan.net_range = [net_totals[first_num], net_totals[last_num]]    
                plan.net_ext_range = [net_ext_totals[first_num], net_ext_totals[last_num]]

                # Test costs
                tol = plan.total_costs * tol_percent + 0.001
                if max(abs(plan.cost_range[0] - old_cost[0]), abs(plan.cost_range[1] - old_cost[1])) < tol:
                    cost_ben_net[0] = True
                # Test bens
                tol = plan.total_bens * tol_percent + 0.001
                if max(abs(plan.ben_range[0] - old_ben[0]), abs(plan.ben_range[1] - old_ben[1])) < tol:
                    cost_ben_net[1] = True
                # Test net
                tol = plan.net * tol_percent + 0.001
                if max(abs(plan.net_range[0] - old_net[0]), abs(plan.net_range[1] - old_net[1])) < tol:
                    cost_ben_net[2] = True
                # Test net with externalities
                tol = plan.net_w_ext * tol_percent + 0.001
                if max(abs(plan.net_ext_range[0] - old_ext_net[0]), abs(plan.net_ext_range[1] - old_ext_net[1])) < tol:
                    cost_ben_net[3] = True

                old_iters = num_iters
                num_iters = 2 * num_iters

                #print('Iterations:', num_iters/2)
                #print('Costs:', plan.cost_range)
                #print('Benefits:', plan.ben_range)
                #print('Net:', plan.net_range)
                #print(cost_ben_net)
                #input()
            ben_direct_totals.sort()
            ben_indirect_totals.sort()
            fat_num_totals.sort()
            fat_value_totals.sort()
            cost_direct_totals.sort()
            cost_indirect_totals.sort()
            cost_omr_1_totals.sort()
            cost_omr_r_totals.sort()
            res_rec_totals.sort()
            nond_1_totals.sort()
            nond_r_totals.sort()
            ext_p_one_totals.sort()
            ext_p_r_totals.sort()
            ext_n_one_totals.sort()
            ext_n_r_totals.sort()
            irr_totals.sort()
            sir_totals.sort()
            roi_totals.sort()
            nond_roi_totals.sort()
            plan.bens.direct_range = [ben_direct_totals[first_num], ben_direct_totals[last_num]]
            plan.bens.indirect_range = [ben_indirect_totals[first_num], ben_indirect_totals[last_num]]
            plan.bens.res_rec_range = [res_rec_totals[first_num], res_rec_totals[last_num]]
            plan.fat.num_range = [fat_num_totals[first_num], fat_num_totals[last_num]]
            plan.fat.value_range = [fat_value_totals[first_num], fat_value_totals[last_num]]
            plan.costs.direct_range = [cost_direct_totals[first_num], cost_direct_totals[last_num]]
            plan.costs.indirect_range = [cost_indirect_totals[first_num], cost_indirect_totals[last_num]]
            plan.costs.omr_one_range = [cost_omr_1_totals[first_num], cost_omr_1_totals[last_num]]
            plan.costs.omr_r_range = [cost_omr_r_totals[first_num], cost_omr_r_totals[last_num]]
            plan.nond_bens.one_range = [nond_1_totals[first_num], nond_1_totals[last_num]]
            plan.nond_bens.r_range = [nond_r_totals[first_num], nond_r_totals[last_num]]
            plan.exts.one_p_range = [ext_p_one_totals[first_num], ext_p_one_totals[last_num]]
            plan.exts.one_n_range = [ext_n_one_totals[first_num], ext_n_one_totals[last_num]]
            plan.exts.r_p_range = [ext_p_r_totals[first_num], ext_p_r_totals[last_num]]
            plan.exts.r_n_range = [ext_n_r_totals[first_num], ext_n_r_totals[last_num]]
            plan.irr_range = [irr_totals[first_num], irr_totals[last_num]]
            plan.sir_range = [sir_totals[first_num], sir_totals[last_num]]         
            plan.roi_range = [roi_totals[first_num], roi_totals[last_num]]        
            plan.nond_roi_range = [nond_roi_totals[first_num], nond_roi_totals[last_num]] 

            if plan.irr_range[0] < 0:
                plan.irr_range[0] = '---'
            if plan.irr_range[1] < 0:
                plan.irr_range[1] = '---'
            plan.mc_iters = num_iters / 2

    def one_iter(self, my_plan):
        """ Creates one plan that is within the uncertainty bounds of my_plan."""
        def to_pass(dist, some_range):
            """Returns the mid and range to pass the distributions given dist and some_range."""
            if dist == "none":
                mid = some_range[0]
                new_range = [0, 0, 0, 0, 0, 0]
            elif dist == "discrete":
                mid = max(some_range[0:2])
                new_range = list(some_range)
            elif dist == "gauss":
                mid = some_range[0]
                new_range = [some_range[1], 0, 0, 0, 0, 0]
            else:
                mid = some_range[0]
                new_range = [some_range[1], some_range[2], 0, 0, 0, 0]
            return [mid, new_range]
        dist_dict = {'tri':triDistInv, 'rect':uniDistInv, 'none':none_dist, 'discrete':discrete_dist_inv, 'gauss':gauss_dist_inv}
        new_recurr = dist_dict[my_plan.recurr_dist](np.random.uniform(), to_pass(my_plan.recurr_dist, my_plan.recurr_range)[0], to_pass(my_plan.recurr_dist, my_plan.recurr_range)[1])
        new_mag = dist_dict[my_plan.mag_dist](np.random.uniform(), to_pass(my_plan.mag_dist, my_plan.mag_range)[0], to_pass(my_plan.mag_dist, my_plan.mag_range)[1])
        delta_plan = Plan(my_plan.id_assign, my_plan.name, [my_plan.recurr_dist, [new_recurr, my_plan.recurr_range]],
                          [my_plan.mag_dist, [new_mag, my_plan.mag_range]], self.discount_rate, self.horizon,
                          self.stat_life, self.parties)
        delta_plan.bens = delta_plan.bens.one_iter(my_plan.bens.indiv)
        delta_plan.exts = delta_plan.exts.one_iter(my_plan.exts.indiv)
        delta_plan.costs = delta_plan.costs.one_iter(my_plan.costs.indiv)
        delta_plan.fat.update(my_plan.fat.averted, my_plan.fat.desc)
        delta_plan.nond_bens = delta_plan.nond_bens.one_iter(my_plan.nond_bens.indiv)
        return delta_plan

class Plan():
    """ Has all of the aspects of the plan. """
    tol = 1E-7
    def __init__(self, plan_id, plan_name, disaster_recurrence, disaster_magnitude,
                 discount_rate, horizon, stat_life, parties):
        # Basic Information
        self.id_assign = int(plan_id)
        self.name = plan_name
        self.recurr_dist = disaster_recurrence[0]
        self.recurr_range = disaster_recurrence[1]
        if self.recurr_dist == "none":
            self.recurr_uncert = [0]
            self.recurrence = disaster_recurrence[1][0]
        elif self.recurr_dist == "gauss":
            self.recurr_uncert = disaster_recurrence[1][1]
            self.recurrence = disaster_recurrence[1][0]
        elif self.recurr_dist == "discrete":
            self.recurr_uncert = list(disaster_recurrence)
            self.recurrence = max([disaster_recurrence[1][0], disaster_recurrence[1][2], disaster_recurrence[1][3]])
        else:
            self.recurr_uncert = list(disaster_recurrence[1])
            self.recurrence = [disaster_recurrence[1][1]]
        self.mag_dist = disaster_magnitude[0]
        self.mag_range = disaster_magnitude[1]
        if self.mag_dist == "none":
            self.mag_uncert = [0]
            self.magnitude = [disaster_magnitude[1][0]]
        elif self.mag_dist == "gauss":
            self.mag_uncert = disaster_magnitude[1][1]
            self.magnitude = [disaster_magnitude[1][0]]
        elif self.mag_dist == "discrete":
            self.mag_uncert = list(disaster_magnitude)
            self.magnitude = max([disaster_magnitude[1][0], disaster_magnitude[1][2], disaster_magnitude[1][3]])
        else:
            self.mag_uncert = list(disaster_magnitude[1])
            self.magnitude = disaster_magnitude[1][1]
        self.horizon = horizon

        # Specific pieces
        self.costs = Costs(discount_rate, horizon)
        self.exts = Externalities(discount_rate, horizon, parties)
        self.bens = Benefits(self.recurrence, discount_rate, horizon)
        self.fat = Fatalities(self.recurrence, discount_rate, stat_life, horizon)
        self.nond_bens = NonDBens(discount_rate, horizon)

        self.total_bens = 0
        self.total_costs = 0
        self.net = 0
        self.net_w_ext = self.net

        self.annual_cash_flows = [0 for x in range(int(horizon) + 1)]
        self.annual_non_disaster_cash_flows = [0 for x in range(int(horizon) + 1)]
        self.ben_range = [0, 0]
        self.cost_range = [0, 0]
        self.net_range = [0, 0]
        self.net_ext_range = [0, 0]


    def update(self, plan_id, plan_name, disaster_recurrence, disaster_magnitude,
                 discount_rate, horizon, stat_life):
        # Basic Information
        self.id_assign = int(plan_id)
        self.name = plan_name
        self.recurr_dist = disaster_recurrence[0]
        self.recurr_range = disaster_recurrence[1]
        if self.recurr_dist == "none":
            self.recurr_uncert = [0]
            self.recurrence = disaster_recurrence[1][0]
        elif self.recurr_dist == "gauss":
            self.recurr_uncert = disaster_recurrence[1][1]
            self.recurrence = disaster_recurrence[1][0]
        elif self.recurr_dist == "discrete":
            self.recurr_uncert = list(disaster_recurrence[1])
            self.recurrence = max([disaster_recurrence[1][0], disaster_recurrence[1][2], disaster_recurrence[1][3]])
        else:
            self.recurr_uncert = list(disaster_recurrence[1])
            self.recurrence = disaster_recurrence[1][1]
        self.mag_dist = disaster_magnitude[0]
        self.mag_range = disaster_magnitude[1]
        if self.mag_dist == "none":
            self.mag_uncert = [0]
            self.magnitude = disaster_magnitude[1][0]
        elif self.mag_dist == "gauss":
            self.mag_uncert = disaster_magnitude[1][1]
            self.magnitude = disaster_magnitude[1][0]
        elif self.mag_dist == "discrete":
            self.mag_uncert = list(disaster_magnitude[1])
            self.magnitude = max([disaster_magnitude[1][0], disaster_magnitude[1][2], disaster_magnitude[1][3]])
        else:
            self.mag_uncert = list(disaster_magnitude[1])
            self.magnitude = disaster_magnitude[1][1]
        self.horizon = horizon

        # Specific pieces
        self.costs.discount_rate = discount_rate
        self.costs.horizon = horizon
        self.exts.discount_rate = discount_rate
        self.exts.horizon = horizon
        self.bens.recurrence = self.recurrence
        self.bens.discount_rate = discount_rate
        self.bens.horizon = horizon
        self.fat.recurrence = self.recurrence
        self.fat.discount_rate = discount_rate
        self.fat.horizon = horizon
        self.nond_bens.discount_rate = discount_rate
        self.nond_bens.horizon = horizon

    def new_plan(self, save_list):
        """ Builds a plan from a list of pieces from the save file."""
        for line in save_list:
            if line[1] == "Costs":
                self.costs.new_cost(list(line[2:]))
            elif line[1] == "Benefits":
                self.bens.new_ben(list(line[2:]))
            elif line[1] == "Externalities":
                self.exts.new_ext(list(line[2:]))
            elif line[1] == "Fatalities":
                self.fat.update(line[2], line[3:])
            elif line[1] == "Non-Disaster Benefits":
                self.nond_bens.new_ben(list(line[2:]))

    def save_plan(self, new_file):
        """ Saves a plan into the file."""
        new_file.write('Plan ' + str(self.id_assign) + ',' + self.name + ',')
        new_file.write(self.recurr_dist + ',')
        if self.recurr_dist in {'discrete', 'rect', 'tri'}:
            to_write = []
            for item in self.recurr_uncert:
                to_write.append(item)
            while len(to_write) < 6:
                to_write.append(0)
            for item in to_write:
                new_file.write(str(item) + ',')
        else:
            to_write = [self.recurrence]
            for item in self.recurr_uncert:
                to_write.append(item)
            while len(to_write) < 6:
                to_write.append(0)
            for item in to_write:
                new_file.write(str(item) + ',')
        new_file.write(self.mag_dist + ',')
        if self.mag_dist in {'discrete', 'rect', 'tri'}:
            to_write = []
            for item in self.mag_uncert:
                to_write.append(item)
            while len(to_write) < 6:
                to_write.append(0)
            for item in to_write:
                new_file.write(str(item) + ',')
        else:
            to_write = [self.magnitude]
            for item in self.mag_uncert:
                to_write.append(item)
            while len(to_write) < 6:
                to_write.append(0)
            for item in to_write:
                new_file.write(str(item) + ',')
        new_file.write('\n')
        for cost in self.costs.indiv:
            new_file.write(',Costs,' + cost.title + ',' + cost.cost_type + ',' + cost.omr_type + ',')
            for time in cost.times:
                new_file.write(str(time) + ',')
            new_file.write(str(cost.amount) + ',' + str(cost.desc) + '\n')
            new_file.write(',Costs,Uncertainty,' + cost.dist)
            for value in cost.range:
                if value != '<insert uncertainty>':
                    new_file.write(',' + str(value))
                else:
                    new_file.write(',0')
            new_file.write('\n')
        for ben in self.bens.indiv:
            new_file.write(',Benefits,' + ben.title + ',' + ben.ben_type + ',')
            new_file.write(str(ben.amount) + ',' + str(ben.desc) + '\n')
            new_file.write(',Benefits,Uncertainty,' + ben.dist)
            for value in ben.range:
                if value != '<insert uncertainty>':
                    new_file.write(',' + str(value))
                else:
                    new_file.write(',0')
            new_file.write('\n')
        for ext in self.exts.indiv:
            new_file.write(',Externalities,' + ext.title + ',' + ext.ext_type)
            for entry in ext.times:
                new_file.write(',' + str(entry))
            new_file.write(',' + str(ext.amount) + ',' + str(ext.desc) + '\n')
            new_file.write(',Externalities,')
            if ext.pm == '+':
                new_file.write('positive,')
            else:
                new_file.write('negative,')
            new_file.write(ext.third_party + '\n')
            new_file.write(',Externalities,Uncertainty,' + ben.dist)
            for value in ext.range:
                if value != '<insert uncertainty>':
                    new_file.write(',' + str(value))
                else:
                    new_file.write(',0')
            new_file.write('\n')
        for nond_ben in self.nond_bens.indiv:
            new_file.write(',Non-Disaster Benefits,' + nond_ben.title + ',' + nond_ben.ben_type + ',')
            for item in nond_ben.times:
                new_file.write(str(item) + ',')
            new_file.write(str(nond_ben.amount) + ',' + str(nond_ben.desc) + '\n')
            new_file.write(',Non-Disaster Benefits,Uncertainty,' + nond_ben.dist)
            for value in nond_ben.range:
                if value != '<insert uncertainty>':
                    new_file.write(',' + str(value))
                else:
                    new_file.write(',0')
            new_file.write('\n')
        new_file.write(',Fatalities,' + str(self.fat.averted) + ',' + str(self.fat.desc) + '\n')

    def sum_it(self, horizon):
        """ Sums up all of the individual pieces. """
        self.total_bens = 0
        self.total_costs = 0
        self.net = 0
        horizon = int(horizon)
        # Note: Fatilites does all summing every time it is updated.
        self.costs.make_sum()
        self.bens.make_sum()
        self.exts.make_sum()
        self.nond_bens.make_sum()

        rec_list = []
        ot_list = []
        for ben in self.nond_bens.indiv:
            if ben.ben_type == "recurring":
                rec_list.append(ben)
            else:
                ot_list.append(ben)
        for cost in self.costs.indiv:
            if cost.omr_type == "recurring":
                rec_list.append(cost)
            else:
                ot_list.append(cost)
        time_series = []
        for item in rec_list:
            start = float(item.times[0])
            rate = float(item.times[1])
            if isinstance(item, NonDBenefit):
                amount = float(item.amount)
            else:
                amount = -float(item.amount)
            i = 0
            while start + rate * i <= horizon + self.tol:
                time_series.append([start + rate * i, amount])
                i += 1
                if rate <= self.tol:
                    break
        for item in ot_list:
            if isinstance(item, NonDBenefit):
                amount = float(item.amount)
            else:
                amount = -float(item.amount)
            time_series.append([float(item.times[0]), amount])

        time_series.sort(key=lambda x: x[0])
        prev = -1
        for i in range(len(time_series)):
            if abs(time_series[i][0]-prev) <= self.tol:
                time_series[i][1] += float(time_series[i-1][1])
                time_series[i-1][0] = -1
            else:
                prev = time_series[i][0]

        self.annual_cash_flows = []
        for item in time_series:
            if item[0] != -1:
                self.annual_cash_flows.append(item)

        self.total_bens = self.bens.total + self.fat.stat_value_averted + self.nond_bens.total
        self.total_costs = self.costs.total
        self.net = self.total_bens - self.total_costs
        self.net_w_ext = self.total_bens + self.exts.total_p - self.total_costs - self.exts.total_n

    def sir(self):
        """Equation for the Savings-to-Investment Ratio"""
        up_front = self.costs.d_sum + self.costs.i_sum

        if up_front == 0:
            return 0

        return self.net / up_front

    def irr(self):
        """Equation for the Internal Rate of Return"""
        # === Calls the function so that self.up_front is calculated
        #self.sir(plan_num)

        #annual_cost = (self.cost.total - self.up_front) / float(self.horizon)
        #annual_savings = self.ben.total / float(self.horizon)

        #irr_list = [annual_savings - annual_cost] * (int(self.horizon) + 1)
        #irr_list[0] = -(self.up_front)

        cash_flows = self.annual_cash_flows
        ben_list = [self.bens.d_sum_no_discount, self.bens.i_sum_no_discount,
                    self.bens.r_sum_no_discount]

        try:
            the_irr = irr_for_all(cash_flows, self.horizon, self.recurrence, ben_list,
                                  self.fat.stat_life, self.fat.averted)
        except ValueError:
            print('ValueError')
            return 'No Valid IRR'
        except OverflowError:
            print('OverflowError')
            return 'No Valid IRR'
        if the_irr == 0.5:
            return "---"

        return the_irr * 100


    def roi(self):
        """Equation for the Return on Investment"""
        if self.total_bens == 0:
            return 0
        elif self.total_costs == 0:
            return 0
        annual_savings = self.total_bens / float(self.horizon)
        simple_payback = self.total_costs / annual_savings
        return (1 / simple_payback) * 100

    def non_d_roi(self):
        """Equation for the Return on Investment (without any chance of disaster occurring)"""
        non_d_ben_total = self.nond_bens.r_sum + self.nond_bens.one_sum

        if non_d_ben_total == 0:
            return 0
        elif self.total_costs == 0:
            return 0
        annual_savings = non_d_ben_total / float(self.horizon)
        simple_payback = self.total_costs / annual_savings
        return (1 / simple_payback) * 100

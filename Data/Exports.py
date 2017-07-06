""" The exporting capabilities.
    Author: Shannon Grubb
            shannon.grubb@nist.gov
    2017-07
"""

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from tkinter import filedialog


def csv_export(sim):
    """ Exports the results as a csv without uncertainty."""
    my_formats = [('Comma Separated Value', '*.csv'),]
    file_name = filedialog.asksaveasfilename(filetypes=my_formats, title="Save the file as...")
    if file_name[-4:] != '.csv':
        file_name = file_name + '.csv'
    new_file = open(file_name, 'w')
    new_file.write('Outputs of Economic Evaluation: [' + sim.title + ']\n')
    new_file.write('Number of Alternatives: ' + str(sim.num_plans-1) + '\n')
    new_file.write('Planning Horizon: ' + str(sim.horizon) + ' years' + '\n')
    new_file.write('Discount Rate: ' + str(sim.discount_rate) + '%' + '\n')
    new_file.write('Disaster Rate: Every ' + str(sim.get_disaster_rate()[0]) + ' years\n')
    new_file.write('Disaster Magnitude: ' + str(sim.get_disaster_magnitude()[0])
                   + '% of build cost' + '\n')
    new_file.write('Risk Preference: ' + str(sim.risk_pref) + '\n')
    new_file.write('Statistical Value of a Life: '+'${:.0f}'.format(float(sim.stat_life))+'\n')
    new_file.write(',Base Case')
    for i in range(1, sim.num_plans):
        new_file.write(',Alternative ' + str(i))
    new_file.write('\n,Base')
    for i in range(1, sim.num_plans):
        new_file.write(',' + sim.plan_list[i].name)
    new_file.write('\nBenefits\nDisaster Economic Benefits\nResponse and Recovery Costs')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.bens.r_sum))
    new_file.write('\nDirect Loss Reduction')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.bens.d_sum))
    new_file.write('\nIndirect Losses')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.bens.i_sum))
    new_file.write('\nDisaster Non-Market Benefits\nValue of Statistical Lives Saved')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.fat.stat_value_averted))
    new_file.write('\nNumber of Statistical Lives Saved')
    for plan in sim.plan_list:
        new_file.write(',' + str(plan.fat.stat_averted))
    new_file.write('\nNon-disaster Related Benefits\nOne-Time')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.nond_bens.one_sum))
    new_file.write('\nRecurring\n')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.nond_bens.r_sum))
    new_file.write('\nCosts\nDirect Costs')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.costs.d_sum))
    new_file.write('\nIndirect Costs')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.costs.i_sum))
    new_file.write('\nOMR\nOne-Time')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.costs.omr_1_sum))
    new_file.write('\nRecurring')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.costs.omr_r_sum))
    new_file.write('\nExternalities\nPositive\nOne-Time')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.exts.one_sum_p))
    new_file.write('\nRecurring')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.exts.r_sum_p))
    new_file.write('\nNegative\nOne-Time')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.exts.one_sum_n))
    new_file.write('\nRecurring')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.exts.r_sum_n))
    new_file.write('\nTotal: Present Expected Value\nBenefits')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.total_bens))
    new_file.write('\nCosts')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.total_costs))
    new_file.write('\nNet')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.net))
    any_ext = False
    for plan in sim.plan_list:
        if len(plan.exts.indiv) > 0:
            any_ext = True
    if any_ext:
        new_file.write('\nNet with Externalities')
        for plan in sim.plan_list:
            new_file.write(',$' + str(plan.net_w_ext))
    new_file.write('\nSavings-to-Investment Ratio')
    for plan in sim.plan_list:
        new_file.write(',' + str(plan.sir()))
    new_file.write('\nInternal Rate of Return (%)')
    for plan in sim.plan_list:
        new_file.write(',' + str(plan.irr()))
    new_file.write('\nReturn on Investment (%)')
    for plan in sim.plan_list:
        new_file.write(',' + str(plan.roi()))
    new_file.write('\nNon-Disaster ROI (%)')
    for plan in sim.plan_list:
        new_file.write(',' + str(plan.non_d_roi()))
    new_file.close()

def csv_export_uncert(sim):
    """ Exports the results with uncertainty as a csv file."""
    my_formats = [('Comma Separated Value', '*.csv'),]
    file_name = filedialog.asksaveasfilename(filetypes=my_formats, title="Save the file as...")
    if file_name[-4:] != '.csv':
        file_name = file_name + '.csv'
    new_file = open(file_name, 'w')
    new_file.write('Outputs of Economic Evaluation: [' + sim.title + ']\n')
    new_file.write('NOTE: All bounds on uncertainties are given with a ' + str(sim.confidence)
                   + '% confidence interval. The number of runs was determined with a '
                   + str(sim.tolerance) + '% tolerance.\n')
    for plan in sim.plan_list:
        new_file.write('For ' + plan.name + ' (Alternative ' + str(plan.id_assign) + ') '
                       + str(plan.mc_iters) + ' Monte-Carlo simulations were run.\n')
    new_file.write('The random number seed for these runs was ' + str(sim.seed) + '.\n')
    new_file.write('Number of Alternatives: ' + str(sim.num_plans-1) + '\n')
    new_file.write('Planning Horizon: ' + str(sim.horizon) + ' years' + '\n')
    new_file.write('Discount Rate: ' + str(sim.discount_rate) + '%' + '\n')
    new_file.write('Disaster Rate: Every ' + str(sim.get_disaster_rate()[0]) + ' years\n')
    new_file.write('Disaster Magnitude: ' + str(sim.get_disaster_magnitude()[0])
                   + '% of build cost' + '\n')
    new_file.write('Risk Preference: ' + str(sim.risk_pref) + '\n')
    new_file.write('Statistical Value of a Life: '+'${:.0f}'.format(float(sim.stat_life))+'\n')

    new_file.write(',Base Case,Lower Bound,Upper Bound')
    for i in range(1, sim.num_plans):
        new_file.write(',Alternative ' + str(i) + ',Lower Bound,Upper Bound')
    new_file.write('\n,Base')
    for i in range(1, sim.num_plans):
        new_file.write(',,,' + sim.plan_list[i].name)
    new_file.write('\nBenefits\nDisaster Economic Benefits\nResponse and Recovery Costs')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.bens.r_sum))
        new_file.write(',$' + str(plan.bens.res_rec_range[0])
                       + ',$' + str(plan.bens.res_rec_range[1]))
    new_file.write('\nDirect Loss Reduction')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.bens.d_sum))
        new_file.write(',$' + str(plan.bens.direct_range[0])
                       + ',$' + str(plan.bens.direct_range[1]))
    new_file.write('\nIndirect Losses')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.bens.i_sum))
        new_file.write(',$' + str(plan.bens.indirect_range[0])
                       + ',$' + str(plan.bens.indirect_range[1]))
    new_file.write('\nDisaster Non-Market Benefits\nValue of Statistical Lives Saved')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.fat.stat_value_averted))
        new_file.write(',$' + str(plan.fat.value_range[0])
                       + ',$' + str(plan.fat.value_range[1]))
    new_file.write('\nNumber of Statistical Lives Saved')
    for plan in sim.plan_list:
        new_file.write(',' + str(plan.fat.stat_averted))
        new_file.write(',' + str(plan.fat.num_range[0]) + ',' + str(plan.fat.num_range[1]))
    new_file.write('\nNon-disaster Related Benefits\nOne-Time')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.nond_bens.one_sum))
        new_file.write(',$' + str(plan.nond_bens.one_range[0])
                       + ',$' + str(plan.nond_bens.one_range[1]))
    new_file.write('\nRecurring')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.nond_bens.r_sum))
        new_file.write(',$' + str(plan.nond_bens.r_range[0])
                       + ',$' + str(plan.nond_bens.r_range[1]))
    new_file.write('\nCosts\nDirect Costs')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.costs.d_sum))
        new_file.write(',$' + str(plan.costs.direct_range[0])
                       + ',$' + str(plan.costs.direct_range[1]))
    new_file.write('\nIndirect Costs')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.costs.i_sum))
        new_file.write(',$' + str(plan.costs.indirect_range[0])
                       + ',$' + str(plan.costs.indirect_range[1]))
    new_file.write('\nOMR\nOne-Time')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.costs.omr_1_sum))
        new_file.write(',$' + str(plan.costs.omr_one_range[0])
                       + ',$' + str(plan.costs.omr_one_range[1]))
    new_file.write('\nRecurring')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.costs.omr_r_sum))
        new_file.write(',$' + str(plan.costs.omr_r_range[0])
                       + ',$' + str(plan.costs.omr_r_range[1]))
    new_file.write('\nExternalities\nPositive\nOne-Time')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.exts.one_sum_p))
        new_file.write(',$' + str(plan.exts.one_p_range[0])
                       + ',$' + str(plan.exts.one_p_range[1]))
    new_file.write('\nRecurring')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.exts.r_sum_p))
        new_file.write(',$' + str(plan.exts.r_p_range[0]) + ',$' + str(plan.exts.r_p_range[1]))
    new_file.write('\nNegative\nOne-Time')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.exts.one_sum_n))
        new_file.write(',$' + str(plan.exts.one_n_range[0])
                       + ',$' + str(plan.exts.one_n_range[1]))
    new_file.write('\nRecurring')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.exts.r_sum_n))
        new_file.write(',$' + str(plan.exts.r_n_range[0]) + ',$' + str(plan.exts.r_n_range[1]))
    new_file.write('\nTotal: Present Expected Value\nBenefits')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.total_bens))
        new_file.write(',$' + str(plan.ben_range[0]) + ',$' + str(plan.ben_range[1]))
    new_file.write('\nCosts')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.total_costs))
        new_file.write(',$' + str(plan.cost_range[0]) + ',$' + str(plan.cost_range[1]))
    new_file.write('\nNet')
    for plan in sim.plan_list:
        new_file.write(',$' + str(plan.net))
        new_file.write(',$' + str(plan.net_range[0]) + ',$' + str(plan.net_range[1]))
    any_ext = False
    for plan in sim.plan_list:
        if len(plan.exts.indiv) > 0:
            any_ext = True
    if any_ext:
        new_file.write('\nNet with Externalities')
        for plan in sim.plan_list:
            new_file.write(',$' + str(plan.net_w_ext))
            new_file.write(',$' + str(plan.net_ext_range[0])
                           + ',$' + str(plan.net_ext_range[1]))
    new_file.write('\nSavings-to-Investment Ratio')
    for plan in sim.plan_list:
        new_file.write(',' + str(plan.sir()))
        new_file.write(',' + str(plan.sir_range[0]) + ',' + str(plan.sir_range[1]))
    new_file.write('\nInternal Rate of Return (%)')
    for plan in sim.plan_list:
        new_file.write(',' + str(plan.irr()))
        new_file.write(',' + str(plan.irr_range[0]) + ',' + str(plan.irr_range[1]))
    new_file.write('\nReturn on Investment (%)')
    for plan in sim.plan_list:
        new_file.write(',' + str(plan.roi()))
        new_file.write(',' + str(plan.roi_range[0]) + ',' + str(plan.roi_range[1]))
    new_file.write('\nNon-Disaster ROI (%)')
    for plan in sim.plan_list:
        new_file.write(',' + str(plan.non_d_roi()))
        new_file.write(',' + str(plan.nond_roi_range[0]) + ',' + str(plan.nond_roi_range[1]))
    new_file.close()

def ben_table(sim, plan, table):
    """ Constructs the benefits table."""
    i = 2
    table.cell(0, 0).text = 'Title'
    table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 1).text = 'Amount ($)'
    table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 2).text = 'Effective Present Value ($)'
    table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(1, 0).text = 'Response and Recovery Cost Reductions'
    table.cell(1, 2).text = '{:,.0f}'.format(plan.bens.r_sum)
    table.cell(1, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for ben in plan.bens.indiv:
        if ben.ben_type == "res-rec":
            table.cell(i, 0).text = ben.title
            table.cell(i, 1).text = '{:,.0f}'.format(ben.amount)
            table.cell(i, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(i, 2).text = '{:,.0f}'.format(plan.bens.on_dis_occ(ben.amount, plan.horizon,
                                                                          plan.recurrence,
                                                                          sim.discount_rate))
            table.cell(i, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            i += 1
    table.cell(i, 0).text = 'Direct Losses Prevented'
    table.cell(i, 2).text = '{:,.0f}'.format(plan.bens.d_sum)
    table.cell(i, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    i += 1
    for ben in plan.bens.indiv:
        if ben.ben_type == "direct":
            table.cell(i, 0).text = ben.title
            table.cell(i, 1).text = '{:,.0f}'.format(ben.amount)
            table.cell(i, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(i, 2).text = '{:,.0f}'.format(plan.bens.on_dis_occ(ben.amount, plan.horizon,
                                                                          plan.recurrence,
                                                                          sim.discount_rate))
            table.cell(i, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            i += 1
    table.cell(i, 0).text = 'Indirect Losses Prevented'
    table.cell(i, 2).text = '{:,.0f}'.format(plan.bens.i_sum)
    table.cell(i, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    i += 1
    for ben in plan.bens.indiv:
        if ben.ben_type == "indirect":
            table.cell(i, 0).text = ben.title
            table.cell(i, 1).text = '{:,.0f}'.format(ben.amount)
            table.cell(i, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(i, 2).text = '{:,.0f}'.format(plan.bens.on_dis_occ(ben.amount, plan.horizon,
                                                                          plan.recurrence,
                                                                          sim.discount_rate))
            table.cell(i, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            i += 1
    table.cell(i, 0).text = 'Total'
    table.cell(i, 2).text = '{:,.0f}'.format(plan.bens.total)
    table.cell(i, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    return table

def nond_ben_table(sim, plan, table):
    """ Constructs the NonDBens table."""
    i = 2
    # Header
    table.cell(0, 0).text = 'Title'
    table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 1).text = 'Start Year'
    table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 2).text = 'Recurrence (Years)'
    table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 3).text = 'Amount ($)'
    table.cell(0, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 4).text = 'Effective Present Value ($)'
    table.cell(0, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(1, 0).text = 'One Time Cost Reductions'
    table.cell(1, 4).text = '{:,.0f}'.format(plan.nond_bens.one_sum)
    table.cell(1, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for ben in plan.nond_bens.indiv:
        if ben.ben_type == "one-time":
            table.cell(i, 0).text = ben.title
            table.cell(i, 1).text = str(ben.times[0])
            table.cell(i, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(i, 2).text = 'N/A'
            table.cell(i, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(i, 3).text = '{:,.0f}'.format(float(ben.amount))
            table.cell(i, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(i, 4).text = '{:,.0f}'.format(plan.nond_bens.calc_one_time(ben.amount,
                                                                                  ben.times[0]))
            table.cell(i, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            i += 1
    table.cell(i, 0).text = 'Recurring Cost Reductions'
    table.cell(i, 4).text = '{:,.0f}'.format(plan.nond_bens.r_sum)
    table.cell(i, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    i += 1
    for ben in plan.nond_bens.indiv:
        if ben.ben_type == "recurring":
            table.cell(i, 0).text = ben.title
            table.cell(i, 1).text = str(ben.times[0])
            table.cell(i, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(i, 2).text = str(ben.times[1])
            table.cell(i, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(i, 3).text = '{:,.0f}'.format(float(ben.amount))
            table.cell(i, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(i, 4).text = '{:,.0f}'.format(plan.nond_bens.calc_recur(ben.amount,
                                                                               ben.times[0],
                                                                               ben.times[1]))
            table.cell(i, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            i += 1
    table.cell(i, 0).text = 'Total'
    table.cell(i, 4).text = '{:,.0f}'.format(plan.nond_bens.total)
    table.cell(i, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    return table

def cost_table(sim, plan, table):
    """ Constructs the cost table."""
    i = 2
    # Header
    table.cell(0, 0).text = 'Title'
    table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 1).text = 'Start Year'
    table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 2).text = 'Recurrence (Years)'
    table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 3).text = 'Amount ($)'
    table.cell(0, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 4).text = 'Effective Present Value ($)'
    table.cell(0, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(1, 0).text = 'Direct Costs'
    table.cell(1, 4).text = '{:,.0f}'.format(plan.costs.d_sum)
    table.cell(1, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for cost in plan.costs.indiv:
        if cost.cost_type == "direct":
            table.cell(i, 0).text = cost.title
            table.cell(i, 1).text = 'Start-Up'
            table.cell(i, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(i, 2).text = 'N/A'
            table.cell(i, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(i, 3).text = '{:,.0f}'.format(cost.amount)
            table.cell(i, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(i, 4).text = '{:,.0f}'.format(plan.costs.calc_one_time(cost.amount,
                                                                              cost.times[0]))
            table.cell(i, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            i += 1
    table.cell(i, 0).text = 'Indirect Costs'
    table.cell(i, 4).text = '{:,.0f}'.format(plan.costs.i_sum)
    table.cell(i, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    i += 1
    for cost in plan.costs.indiv:
        if cost.cost_type == "indirect":
            table.cell(i, 0).text = cost.title
            table.cell(i, 1).text = 'Start-Up'
            table.cell(i, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(i, 2).text = 'N/A'
            table.cell(i, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(i, 3).text = '{:,.0f}'.format(cost.amount)
            table.cell(i, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(i, 4).text = '{:,.0f}'.format(plan.costs.calc_one_time(cost.amount,
                                                                              cost.times[0]))
            table.cell(i, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            i += 1
    table.cell(i, 0).text = 'OMR Costs: One-Time'
    table.cell(i, 4).text = '{:,.0f}'.format(plan.costs.omr_1_sum)
    table.cell(i, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    i += 1
    for cost in plan.costs.indiv:
        if (cost.cost_type == "omr") & (cost.omr_type == "one-time"):
            table.cell(i, 0).text = cost.title
            table.cell(i, 1).text = str(cost.times[0])
            table.cell(i, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(i, 2).text = 'N/A'
            table.cell(i, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(i, 3).text = '{:,.0f}'.format(cost.amount)
            table.cell(i, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(i, 4).text = '{:,.0f}'.format(plan.costs.calc_one_time(cost.amount,
                                                                              float(cost.times[0])))
            table.cell(i, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            i += 1
    table.cell(i, 0).text = 'OMR Costs: Recurring'
    table.cell(i, 4).text = '{:,.0f}'.format(plan.costs.omr_r_sum)
    table.cell(i, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    i += 1
    for cost in plan.costs.indiv:
        if (cost.cost_type == "omr") & (cost.omr_type == "recurring"):
            table.cell(i, 0).text = cost.title
            table.cell(i, 1).text = str(cost.times[0])
            table.cell(i, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(i, 2).text = str(cost.times[1])
            table.cell(i, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(i, 3).text = '{:,.0f}'.format(cost.amount)
            table.cell(i, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(i, 4).text = '{:,.0f}'.format(plan.costs.calc_recur(cost.amount,
                                                                           cost.times[0],
                                                                           cost.times[1]))
            table.cell(i, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            i += 1
    table.cell(i, 0).text = 'Total'
    table.cell(i, 4).text = '{:,.0f}'.format(plan.costs.total)
    table.cell(i, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    return table

def ext_table(sim, plan, table):
    """ Creates Externalities table."""
    i = 2
    # Header
    table.cell(0, 0).text = 'Title'
    table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 1).text = 'Start Year'
    table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 2).text = 'Recurrence (Years)'
    table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 3).text = 'Amount ($)'
    table.cell(0, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 4).text = 'Effective Present Value ($)'
    table.cell(0, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(1, 0).text = 'One Time Positive Externalities'
    table.cell(1, 4).text = '{:,.0f}'.format(plan.exts.one_sum_p)
    table.cell(1, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for ext in plan.exts.indiv:
        if (ext.ext_type == "one-time") & (ext.pm == "+"):
            table.cell(i, 0).text = ext.title
            table.cell(i, 1).text = str(ext.times[0])
            table.cell(i, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(i, 2).text = 'N/A'
            table.cell(i, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(i, 3).text = '{:,.0f}'.format(float(ext.amount))
            table.cell(i, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(i, 4).text = '{:,.0f}'.format(plan.exts.calc_one_time(ext.amount,
                                                                             ext.times[0]))
            table.cell(i, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            i += 1
    table.cell(i, 0).text = 'Recurring Positive Externalities'
    table.cell(i, 4).text = '{:,.0f}'.format(plan.exts.r_sum_p)
    table.cell(i, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    i += 1
    for ext in plan.exts.indiv:
        if (ext.ext_type == "recurring") & (ext.pm == "+"):
            table.cell(i, 0).text = ext.title
            table.cell(i, 1).text = str(ext.times[0])
            table.cell(i, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(i, 2).text = str(ext.times[1])
            table.cell(i, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(i, 3).text = '{:,.0f}'.format(float(ext.amount))
            table.cell(i, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(i, 4).text = '{:,.0f}'.format(plan.exts.calc_recur(ext.amount,
                                                                          ext.times[0],
                                                                          ext.times[1]))
            table.cell(i, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            i += 1
    table.cell(i, 0).text = 'One Time Negative Externalities'
    table.cell(i, 4).text = '{:,.0f}'.format(plan.exts.one_sum_n)
    table.cell(i, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    i += 1
    for ext in plan.exts.indiv:
        if (ext.ext_type == "one-time") & (ext.pm == "-"):
            table.cell(i, 0).text = ext.title
            table.cell(i, 1).text = str(ext.times[0])
            table.cell(i, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(i, 2).text = 'N/A'
            table.cell(i, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(i, 3).text = '{:,.0f}'.format(float(ext.amount))
            table.cell(i, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(i, 4).text = '{:,.0f}'.format(plan.exts.calc_one_time(ext.amount,
                                                                             ext.times[0]))
            table.cell(i, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            i += 1
    table.cell(i, 0).text = 'Recurring Negative Externalities'
    table.cell(i, 4).text = '{:,.0f}'.format(plan.exts.r_sum_n)
    table.cell(i, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    i += 1
    for ext in plan.exts.indiv:
        if (ext.ext_type == "recurring") & (ext.pm == "-"):
            table.cell(i, 0).text = ext.title
            table.cell(i, 1).text = str(ext.times[0])
            table.cell(i, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(i, 2).text = str(ext.times[1])
            table.cell(i, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(i, 3).text = '{:,.0f}'.format(float(ext.amount))
            table.cell(i, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(i, 4).text = '{:,.0f}'.format(plan.exts.calc_recur(ext.amount,
                                                                          ext.times[0],
                                                                          ext.times[1]))
            table.cell(i, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            i += 1
    table.cell(i, 0).text = 'Total'
    table.cell(i, 4).text = '{:,.0f}'.format(plan.exts.total_p - plan.exts.total_n)
    table.cell(i, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    return table

def word_export(sim):
    """ Exports the results (w/o uncertainty) as a Word file."""
    my_formats = [('Microsoft Word Document', '*.docx'),]
    file_name = filedialog.asksaveasfilename(filetypes=my_formats, title="Save the file as...")
    if file_name[-5:] != '.docx':
        file_name = file_name + '.docx'

    any_ext = False
    for plan in sim.plan_list:
        if len(plan.exts.indiv) > 0:
            any_ext = True

    doc = docx.Document()
    for section in doc.sections:
        section.left_margin = 914400
        section.right_margin = 914400

    doc.add_heading('Economic Evaluation Complete Report\n' + sim.title, 0)

    doc.add_heading('Analysis Base Information\n', 1)
    doc.add_paragraph('Number of Alternatives: ' + str(sim.num_plans-1), style='ListBullet')
    doc.add_paragraph('Planning Horizon: ' + str(sim.horizon) + ' years', style='ListBullet')
    doc.add_paragraph('Discount Rate: ' + str(sim.discount_rate) + '%', style='ListBullet')
    doc.add_paragraph()
    doc.add_paragraph('Disaster Rate: Every ' + str(sim.get_disaster_rate()[0]) + ' years',
                      style='ListBullet')
    doc.add_paragraph('Disaster Magnitude: ' + str(sim.get_disaster_magnitude()[0])
                      + '% of build cost', style='ListBullet')
    doc.add_paragraph('Risk Preference: ' + str(sim.risk_pref), style='ListBullet')
    doc.add_paragraph()
    doc.add_paragraph('Statistical Value of a Life: '+'${:.0f}'.format(float(sim.stat_life)),
                      style='ListBullet')

    if any_ext:
        header_list = ['Plan Title', 'Total Benefits ($)', 'Total Costs ($)', 'Net ($)',
                       'Net with externalities ($)', 'SIR (%)', 'IRR (%)', 'ROI (%)',
                       'Non-Disaster ROI (%)']
    else:
        header_list = ['Plan Title', 'Total Benefits ($)', 'Total Costs ($)', 'Net ($)',
                       'SIR (%)', 'IRR (%)', 'ROI (%)', 'Non-Disaster ROI (%)']

    doc.add_heading('Summary\n', 1)
    sum_tab = doc.add_table(rows=len(sim.plan_list) + 1, cols=len(header_list),
                            style='Light List Accent 1')

    for i in range(len(header_list)):
        sum_tab.cell(0, i).text = header_list[i]
        sum_tab.cell(0, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    sum_index = 1
    for plan in sim.plan_list:
        sum_tab.cell(sum_index, 0).text = plan.name
        sum_tab.cell(sum_index, 1).text = '{:,.0f}'.format(plan.total_bens)
        sum_tab.cell(sum_index, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sum_tab.cell(sum_index, 2).text = '{:,.0f}'.format(plan.total_costs)
        sum_tab.cell(sum_index, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sum_tab.cell(sum_index, 3).text = '{:,.0f}'.format(plan.net)
        sum_tab.cell(sum_index, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if any_ext:
            sum_tab.cell(sum_index, 4).text = '{:,.0f}'.format(plan.net_w_ext)
            sum_tab.cell(sum_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sum_tab.cell(sum_index, 5).text = '{:,.2f}'.format(plan.sir())
            sum_tab.cell(sum_index, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if isinstance(plan.irr(), str):
                sum_tab.cell(sum_index, 6).text = plan.irr()
                sum_tab.cell(sum_index, 6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                sum_tab.cell(sum_index, 6).text = '{:,.2f}'.format(plan.irr())
                sum_tab.cell(sum_index, 6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sum_tab.cell(sum_index, 7).text = '{:,.2f}'.format(plan.roi())
            sum_tab.cell(sum_index, 7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sum_tab.cell(sum_index, 8).text = '{:,.2f}'.format(plan.non_d_roi())
            sum_tab.cell(sum_index, 8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sum_index += 1
        else:
            sum_tab.cell(sum_index, 4).text = '{:,.2f}'.format(plan.sir())
            sum_tab.cell(sum_index, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if isinstance(plan.irr(), str):
                sum_tab.cell(sum_index, 5).text = plan.irr()
                sum_tab.cell(sum_index, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                sum_tab.cell(sum_index, 5).text = '{:,.2f}'.format(plan.irr())
                sum_tab.cell(sum_index, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sum_tab.cell(sum_index, 6).text = '{:,.2f}'.format(plan.roi())
            sum_tab.cell(sum_index, 6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sum_tab.cell(sum_index, 7).text = '{:,.2f}'.format(plan.non_d_roi())
            sum_tab.cell(sum_index, 7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sum_index += 1

    for plan in sim.plan_list:
        doc.add_heading(plan.name, 1)
        doc.add_heading('Alternative ' + str(plan.id_assign), 3)
        doc.add_heading('Fatalities Averted\n', 2)
        doc.add_paragraph('Number of Statistical Lives Saved: '
                          + '{:,.2f}'.format(plan.fat.stat_averted))
        doc.add_paragraph('Value of Statistical Lives Saved: '
                          + '${:,.0f}'.format(plan.fat.stat_value_averted))
        if plan.fat.desc != 'N/A':
            doc.add_paragraph('Description: ' + plan.fat.desc)
        doc.add_heading('Disaster-Related Benefits\n', 2)
        # == BENEFITS
        table = doc.add_table(rows=len(plan.bens.indiv)+5, cols=3,
                              style='Light List Accent 1')
        table = ben_table(sim, plan, table)
        for ben in plan.bens.indiv:
            if ben.desc != 'N/A':
                doc.add_paragraph(ben.title + ': ' + ben.desc)

        # == NON-D BENEFITS
        doc.add_heading('Resilience Dividend\n', 2)
        table = doc.add_table(rows=len(plan.nond_bens.indiv)+4, cols=5,
                              style='Light List Accent 1')
        table = nond_ben_table(sim, plan, table)
        for ben in plan.nond_bens.indiv:
            if ben.desc != 'N/A':
                doc.add_paragraph(ben.title + ': ' + ben.desc)

        # == COSTS
        doc.add_heading('Costs\n', 2)
        table = doc.add_table(rows=len(plan.costs.indiv)+6, cols=5, style='Light List Accent 1')
        table = cost_table(sim, plan, table)
        for cost in plan.costs.indiv:
            if cost.desc != 'N/A':
                doc.add_paragraph(cost.title + ': ' + cost.desc)

        # == EXTERNALITIES
        doc.add_heading('Externalities\n', 2)
        table = doc.add_table(rows=len(plan.exts.indiv)+6, cols=5,
                              style='Light List Accent 1')
        table = ext_table(sim, plan, table)
        for ext in plan.exts.indiv:
            if ext.desc != 'N/A':
                doc.add_paragraph(ext.title + ': ' + ext.desc)

    doc.save(file_name)

def uncert_string(uncert, values):
    """Creates a nice string given a certain uncertainty type and the associated values."""
    values = list(values)
    try:
        if (uncert in {'tri', 'rect'}) & ((str(values[2]) == str(0)) | ('insert' in values[2])):
            values[2] = values[1]
    except IndexError:
        pass
    if uncert == "none":
        return "N/A"
    elif uncert == "gauss":
        return "Gaussian distribution with variance of " + str(values[0])
    elif uncert == "tri":
        return ("Triangular distribution with a min of " + str(values[0])
                + " and a max of " + str(values[2]))
    elif uncert == "rect":
        return ("Rectangular distribution with a min of " + str(values[0])
                + " and a max of " + str(values[2]))
    elif uncert == "discrete":
        return (str(values[3]) + "% chance of " + str(values[0]) + ", "
                + str(values[4]) + "% chance of " + str(values[1]) + ", and "
                + str(values[5]) + "% chance of " + str(values[2]))

def word_export_uncert(sim):
    """ Exports results as a Word Document with uncertainties."""
    my_formats = [('Microsoft Word Document', '*.docx'),]

    file_name = filedialog.asksaveasfilename(filetypes=my_formats, title="Save the file as...")
    if file_name[-5:] != '.docx':
        file_name = file_name + '.docx'

    any_ext = False
    for plan in sim.plan_list:
        if len(plan.exts.indiv) > 0:
            any_ext = True

    doc = docx.Document()

    if sim.num_plans > 4:
        section = doc.sections[-1]
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height
    for section in doc.sections:
        section.left_margin = 914400
        section.right_margin = 914400


    doc.add_heading('Economic Evaluation Complete Report\n' + sim.title, 0)
    doc.add_paragraph('NOTE: All bounds on uncertainties are given with a '
                      + str(sim.confidence) + '% confidence interval. '
                      + 'The number of runs was determined with a '
                      + str(sim.tolerance) + '% tolerance.')
    for plan in sim.plan_list:
        doc.add_paragraph('For ' + plan.name + ' (Alternative ' + str(plan.id_assign) + ') '
                          + str(plan.mc_iters) + ' Monte-Carlo simulations were run.')
    doc.add_paragraph('The random number seed for these runs was ' + str(sim.seed) + '.')


    doc.add_heading('Analysis Base Information\n', 1)
    doc.add_paragraph('Number of Alternatives: ' + str(sim.num_plans-1), style='ListBullet')
    doc.add_paragraph('Planning Horizon: ' + str(sim.horizon) + ' years', style='ListBullet')
    doc.add_paragraph('Discount Rate: ' + str(sim.discount_rate) + '%', style='ListBullet')
    doc.add_paragraph()
    doc.add_paragraph('Disaster Rate: Every ' + str(sim.get_disaster_rate()[0]) + ' years',
                      style='ListBullet')
    doc.add_paragraph('Uncertainty in Disaster Rate: '
                      + uncert_string(sim.get_disaster_rate()[2], sim.get_disaster_rate()[1]),
                      style='ListBullet')
    doc.add_paragraph('Disaster Magnitude: ' + str(sim.get_disaster_magnitude()[0])
                      + '% of build cost', style='ListBullet')
    doc.add_paragraph('Uncertainty in Disaster Magnitude: '
                      + uncert_string(sim.get_disaster_magnitude()[2],
                                      sim.get_disaster_magnitude()[1]), style='ListBullet')
    doc.add_paragraph('Risk Preference: ' + str(sim.risk_pref), style='ListBullet')
    doc.add_paragraph()
    doc.add_paragraph('Statistical Value of a Life: '+'${:.0f}'.format(float(sim.stat_life)),
                      style='ListBullet')

    if any_ext:
        header_list = ['', 'Total Benefits ($)', '(Lower Bound, Upper Bound) ($)',
                       'Total Costs ($)', '(Lower Bound, Upper Bound) ($)', 'Net ($)',
                       '(Lower Bound, Upper Bound) ($)', 'Net with externalities ($)',
                       '(Lower Bound, Upper Bound) ($)', 'SIR (%)',
                       '(Lower Bound, Upper Bound) (%)', 'IRR (%)',
                       '(Lower Bound, Upper Bound) (%)', 'ROI (%)',
                       '(Lower Bound, Upper Bound) (%)', 'Non-Disaster ROI (%)',
                       '(Lower Bound, Upper Bound) (%)']
    else:
        header_list = ['', 'Total Benefits ($)', '(Lower Bound, Upper Bound) ($)',
                       'Total Costs ($)', '(Lower Bound, Upper Bound) ($)', 'Net ($)',
                       '(Lower Bound, Upper Bound) ($)', 'SIR (%)',
                       '(Lower Bound, Upper Bound) (%)', 'IRR (%)',
                       '(Lower Bound, Upper Bound) (%)', 'ROI (%)',
                       '(Lower Bound, Upper Bound) (%)', 'Non-Disaster ROI (%)',
                       '(Lower Bound, Upper Bound) (%)']

    doc.add_heading('Summary\n', 1)
    sum_tab = doc.add_table(rows=len(header_list), cols=len(sim.plan_list) + 1,
                            style='Light List Accent 1')

    for i in range(len(header_list)):
        sum_tab.cell(i, 0).text = header_list[i]
        sum_tab.cell(i, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    sum_index = 1
    for plan in sim.plan_list:
        sum_tab.cell(0, sum_index).text = plan.name
        sum_tab.cell(0, sum_index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        sum_tab.cell(1, sum_index).text = '{:,.0f}'.format(plan.total_bens)
        sum_tab.cell(1, sum_index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sum_tab.cell(2, sum_index).text = ('(' + '{:,.0f}'.format(plan.ben_range[0]) + '; '
                                           + '{:,.0f}'.format(plan.ben_range[1]) + ')')
        sum_tab.cell(2, sum_index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sum_tab.cell(3, sum_index).text = '{:,.0f}'.format(plan.total_costs)
        sum_tab.cell(3, sum_index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sum_tab.cell(4, sum_index).text = ('(' + '{:,.0f}'.format(plan.cost_range[0]) + '; '
                                           + '{:,.0f}'.format(plan.cost_range[1]) + ')')
        sum_tab.cell(4, sum_index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sum_tab.cell(5, sum_index).text = '{:,.0f}'.format(plan.net)
        sum_tab.cell(5, sum_index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sum_tab.cell(6, sum_index).text = ('(' + '{:,.0f}'.format(plan.net_range[0]) + '; '
                                           + '{:,.0f}'.format(plan.net_range[1]) + ')')
        sum_tab.cell(6, sum_index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if any_ext:
            sum_tab.cell(7, sum_index).text = '{:,.0f}'.format(plan.net_w_ext)
            sum_tab.cell(7, sum_index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sum_tab.cell(8, sum_index).text = ('(' + '{:,.0f}'.format(plan.net_ext_range[0]) + '; '
                                            + '{:,.0f}'.format(plan.net_ext_range[1]) + ')')
            sum_tab.cell(8, sum_index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sum_tab.cell(9, sum_index).text = '{:,.2f}'.format(plan.sir())
            sum_tab.cell(9, sum_index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sum_tab.cell(10, sum_index).text = ('(' + '{:,.0f}'.format(plan.sir_range[0]) + '; '
                                            + '{:,.0f}'.format(plan.sir_range[1]) + ')')
            sum_tab.cell(10, sum_index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if isinstance(plan.irr(), str):
                sum_tab.cell(11, sum_index).text = plan.irr()
                sum_tab.cell(11, sum_index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                sum_tab.cell(12, sum_index).text = ('(' + plan.irr_range[0] + '; '
                                                + plan.irr_range[1] + ')')
                sum_tab.cell(12, sum_index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                sum_tab.cell(11, sum_index).text = '{:,.2f}'.format(plan.irr())
                sum_tab.cell(11, sum_index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                sum_tab.cell(12, sum_index).text = ('(' + '{:,.2f}'.format(plan.irr_range[0]) + '; '
                                                + '{:,.2f}'.format(plan.irr_range[1]) + ')')
                sum_tab.cell(12, sum_index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sum_tab.cell(13, sum_index).text = '{:,.2f}'.format(plan.roi())
            sum_tab.cell(13, sum_index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sum_tab.cell(14, sum_index).text = ('(' + '{:,.2f}'.format(plan.roi_range[0]) + '; '
                                            + '{:,.2f}'.format(plan.roi_range[1]) + ')')
            sum_tab.cell(14, sum_index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sum_tab.cell(15, sum_index).text = '{:,.2f}'.format(plan.non_d_roi())
            sum_tab.cell(15, sum_index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sum_tab.cell(16, sum_index).text = ('(' + '{:,.2f}'.format(plan.nond_roi_range[0]) + '; '
                                            + '{:,.2f}'.format(plan.nond_roi_range[1]) + ')')
            sum_tab.cell(16, sum_index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sum_index += 1
        else:
            sum_tab.cell(7, sum_index).text = '{:,.2f}'.format(plan.sir())
            sum_tab.cell(7, sum_index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sum_tab.cell(8, sum_index).text = ('(' + '{:,.2f}'.format(plan.sir_range[0]) + '; '
                                            + '{:,.2f}'.format(plan.sir_range[1]) + ')')
            sum_tab.cell(8, sum_index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if isinstance(plan.irr(), str):
                sum_tab.cell(9, sum_index).text = plan.irr()
                sum_tab.cell(9, sum_index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                sum_tab.cell(10, sum_index).text = ('(' + '{:,.2f}'.format(plan.irr_range[0]) + '; '
                                                + '{:,.2f}'.format(plan.irr_range[1]) + ')')
                sum_tab.cell(10, sum_index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                sum_tab.cell(9, sum_index).text = '{:,.2f}'.format(plan.irr())
                sum_tab.cell(9, sum_index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                sum_tab.cell(10, sum_index).text = ('(' + '{:,.2f}'.format(plan.irr_range[0]) + '; '
                                                + '{:,.2f}'.format(plan.irr_range[1]) + ')')
                sum_tab.cell(10, sum_index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sum_tab.cell(11, sum_index).text = '{:,.2f}'.format(plan.roi())
            sum_tab.cell(11, sum_index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sum_tab.cell(12, sum_index).text = ('(' + '{:,.2f}'.format(plan.roi_range[0]) + '; '
                                            + '{:,.2f}'.format(plan.roi_range[1]) + ')')
            sum_tab.cell(12, sum_index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sum_tab.cell(13, sum_index).text = '{:,.2f}'.format(plan.non_d_roi())
            sum_tab.cell(13, sum_index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sum_tab.cell(14, sum_index).text = ('(' + '{:,.2f}'.format(plan.nond_roi_range[0]) + '; '
                                            + '{:,.2f}'.format(plan.nond_roi_range[1]) + ')')
            sum_tab.cell(14, sum_index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sum_index += 1

    for plan in sim.plan_list:
        doc.add_heading(plan.name, 1)
        doc.add_heading('Alternative ' + str(plan.id_assign), 3)
        doc.add_heading('Fatalities Averted\n', 2)
        doc.add_paragraph('Number of Statistical Lives Saved: '
                          + '{:,.2f}'.format(plan.fat.stat_averted))
        doc.add_paragraph('Value of Statistical Lives Saved: '
                          + '${:,.0f}'.format(plan.fat.stat_value_averted))
        if plan.fat.desc != 'N/A':
            doc.add_paragraph('Description: ' + plan.fat.desc)
        doc.add_heading('Disaster-Related Benefits\n', 2)
        # == BENEFITS
        table = doc.add_table(rows=len(plan.bens.indiv)+5, cols=3, style='Light List Accent 1')
        table = ben_table(sim, plan, table)
        for ben in plan.bens.indiv:
            if ben.dist != 'none':
                doc.add_paragraph(ben.title + ': ' + uncert_string(ben.dist, ben.range))
            if ben.desc != 'N/A':
                doc.add_paragraph(ben.title + ': ' + ben.desc)

        # == NON-D BENEFITS
        doc.add_heading('Resilience Dividend\n', 2)
        table = doc.add_table(rows=len(plan.nond_bens.indiv) + 4, cols=5,
                              style='Light List Accent 1')
        table = nond_ben_table(sim, plan, table)
        for ben in plan.nond_bens.indiv:
            if ben.dist != 'none':
                doc.add_paragraph(ben.title + ': ' + uncert_string(ben.dist, ben.range))
            if ben.desc != 'N/A':
                doc.add_paragraph(ben.title + ': ' + ben.desc)

        # == COSTS
        doc.add_heading('Costs\n', 2)
        table = doc.add_table(rows=len(plan.costs.indiv)+6, cols=5, style='Light List Accent 1')
        table = cost_table(sim, plan, table)
        for cost in plan.costs.indiv:
            if cost.dist != 'none':
                doc.add_paragraph(cost.title + ': ' + uncert_string(cost.dist, cost.range))
            if cost.desc != 'N/A':
                doc.add_paragraph(cost.title + ': ' + cost.desc)

        # == EXTERNALITIES
        doc.add_heading('Externalities\n', 2)
        table = doc.add_table(rows=len(plan.exts.indiv)+6, cols=5,
                              style='Light List Accent 1')
        table = ext_table(sim, plan, table)
        for ext in plan.exts.indiv:
            if ext.dist != 'none':
                doc.add_paragraph(ext.title + ': ' + uncert_string(ext.dist, ext.range))
            if ext.desc != 'N/A':
                doc.add_paragraph(ext.title + ': ' + ext.desc)

    doc.save(file_name)

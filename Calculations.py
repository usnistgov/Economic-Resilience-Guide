""" Contains all calculations for EconGuide and the class that holds all of the data."""
#File:          Calculations.py
#Author:        Shannon Craig
#Email:         shannon.craig@nist.gov
#Date:          October 2016
#Description:   Interacts with EconGuide.py, focusing on the meat of the resilience calculations

import csv
    # Enables use of csv methods. look online for more info
import math
#import tkinter as tk
from tkinter import filedialog
#from tkinter import ttk
#from NewIRR import irr as calc_irr
from NewIRR import irr_for_all
#import numpy
    # Allows for the calculation of Internal Rate of Return (IRR).
    # This must be downloaded from the internet
import docx
# Allows analysis to be exported to a .docx file.
# This must be downloaded from the internet
from Constants import TAB

from BenefitsClass import Benefits
from CostClass import Costs

class Data():
    """ Contains all of the data for the simulations and does a bulk of the calculations.    """

# ========= Functions that complete meaningful calculations
    def on_dis_occ(self, value):
        """Equation used for Expected Value on disaster occurrence"""
        eqn_lambda = 1 / float(self.disaster_rate)
        k = float(self.discount_rate)/100
        mult = eqn_lambda / math.fabs(1 - math.exp(-k))
        return mult * (1 - math.exp(-k * float(self.horizon))) * value

    def calc_fatalities(self, value):
        """Equation used for Fatalities Aversion Calculation"""
        # == may need additional calculations
        return self.on_dis_occ(float(self.stat_life) * float(value))

    def calc_one_time(self, value, time):
        """Equation used for One-time OMR costs"""
        value = float(value)
        time = float(time)
        return (math.exp(-(float(self.discount_rate) / 100) * time)) * value

    def calc_recur(self, value, start, rate):
        """Equation used for Recurring OMR costs"""
        value = float(value)
        year = float(start)
        rate = float(rate)
        total = 0

        while year <= float(self.horizon):
            total += value * math.exp(-(float(self.discount_rate) / 100) * (year))
            year += rate
        return total

    def sir(self, plan_num):
        """Equation for the Savings-to-Investment Ratio"""
        for i in range(self.num_plans + 1):
            self.up_front.append(self.cost.d_sum[i] + self.cost.i_sum[i])

        if self.up_front[plan_num] == 0:
            return 0

        return self.net[plan_num] / self.up_front[plan_num]

    def irr(self, plan_num):
        """Equation for the Internal Rate of Return"""
        # === Calls the function so that self.up_front is calculated
        #self.sir(plan_num)

        #annual_cost = (self.cost.total[plan_num] - self.up_front[plan_num]) / float(self.horizon)
        #annual_savings = self.ben.total[plan_num] / float(self.horizon)

        #irr_list = [annual_savings - annual_cost] * (int(self.horizon) + 1)
        #irr_list[0] = -(self.up_front[plan_num])

        cash_flows = self.annual_non_disaster_cash_flows[plan_num]
        ben_list = [self.ben.d_sum[plan_num], self.ben.i_sum[plan_num],
                    self.ben.r_sum[plan_num]]

        try:
            the_irr = irr_for_all(cash_flows, self.horizon, self.disaster_rate, ben_list,
                                  self.stat_life, self.Fatalities[plan_num][1])
        except ValueError:
            print('ValueError')
            return 'No Valid IRR'
        except OverflowError:
            print('OverflowError')
            return 'No Valid IRR'
        if the_irr == 0.5:
            return "---"

        return the_irr * 100


    def roi(self, plan_num):
        """Equation for the Return on Investment"""
        if self.ben.total[plan_num] == 0:
            return 0
        elif self.cost.total[plan_num] == 0:
            return 0
        annual_savings = self.ben.total[plan_num] / float(self.horizon)
        simple_payback = self.cost.total[plan_num] / annual_savings
        return (1 / simple_payback) * 100

    def non_d_roi(self, plan_num):
        """Equation for the Return on Investment (without any chance of disaster occurring)"""
        non_d_ben_total = self.recur_non_d_ben_sum[plan_num] + self.one_time_non_d_ben_sum[plan_num]

        if non_d_ben_total == 0:
            return 0
        elif self.cost.total[plan_num] == 0:
            return 0
        annual_savings = non_d_ben_total / float(self.horizon)
        simple_payback = self.cost.total[plan_num] / annual_savings
        return (1 / simple_payback) * 100

    def summer(self):
        """ Sums up all of the values."""
        self.ben.r_sum = [0]*(self.num_plans + 1)
        self.ben.d_sum = [0]*(self.num_plans + 1)
        self.ben.i_sum = [0]*(self.num_plans + 1)
        self.fatalities_sum = [0]*(self.num_plans + 1)
        self.cost.d_sum = [0]*(self.num_plans + 1)
        self.cost.i_sum = [0]*(self.num_plans + 1)
        self.one_time_non_d_ben_sum = [0]*(self.num_plans + 1)
        self.recur_non_d_ben_sum = [0]*(self.num_plans + 1)
        self.cost.omr_1_sum = [0]*(self.num_plans + 1)
        self.cost.omr_r_sum = [0]*(self.num_plans + 1)
        self.cost.total = [0]*(self.num_plans + 1)
        self.ben.total = [0]*(self.num_plans + 1)
        self.net = [0]*(self.num_plans + 1)
        self.annual_non_disaster_cash_flows = [[0]*(int(self.horizon)+1)
                                               for i in range(int(self.num_plans)+1)]

        for i in range(len(self.NonDBen)):
            for j in range(len(self.NonDBen[i])):
                if (self.NonDBen[i][j][0] not in self.nonDBenNames) and self.NonDBen[i][j][0] != "":
                    self.nonDBenNames.append([self.NonDBen[i][j][0]])


        for i in range(self.num_plans + 1):
            # === Response/Recovery Costs, Direct Costs, Indirect Costs Reduction

            # === Fatalities Aversion

            # === Direct Costs, Indirect Costs, One Time/Recurring OMR
            cash_flows[i][0] -= float(self.direct[i][j][1])
            cash_flows[i][0] -= float(self.indirect[i][j][1])
            for k in range(1, len(cash_flows[i])):
                if k % int(self.omr_type[i][j][2]) == 0:
                    cash_flows[i][k] += -float(self.omr[i][j][1])

        # === NonDBens
            self.annual_non_disaster_cash_flows[i][int(self.NonDBen[i][j][4])] += float(self.NonDBen[i][j][1])
            for k in range(1, len(self.annual_non_disaster_cash_flows[i])):
                if k % int(self.NonDBen[i][j][5]) == 0:
                    self.annual_non_disaster_cash_flows[i][k] += float(self.NonDBen[i][j][1])

        # === Totals
        for i in range(self.num_plans + 1):
            self.ben.total[i] += self.on_dis_occ(self.ben.r_sum[i])
            self.ben.total[i] += self.on_dis_occ(self.ben.d_sum[i])
            self.ben.total[i] += self.on_dis_occ(self.ben.i_sum[i])
            self.ben.total[i] += self.fatalities_sum[i]
            self.ben.total[i] += self.one_time_non_d_ben_sum[i] + float(self.recur_non_d_ben_sum[i])
            self.cost.total[i] += float(self.cost.d_sum[i]) + float(self.cost.i_sum[i])
            self.cost.total[i] += float(self.cost.omr_1_sum[i]) + float(self.cost.omr_r_sum[i])
            self.net[i] = self.ben.total[i] - self.cost.total[i]


# ========= Functions that interact with a text file =========

    def save_info(self):
        """Saves current progress
        First converts all information into a 2D array, then writes into .csv file.
        Accessible to all classes"""

        # === Holds the maximum lengths of any list related to each plan
        max_lengths = []
        for i in range(self.num_plans + 1):
            lengths = [len(self.cost.direct[i]), len(self.cost.indirect[i]),
                       len(self.cost.omr[i]), len(self.Externalities[i]),
                       len(self.ben.direct[i]), len(self.ben.indirect[i]),
                       len(self.ben.res_rec[i]), len(self.NonDBen[i])]
            max_lengths.append(max(lengths))

        to_write_list = [["Analysis Title", self.analysis_title],   #0
                         ["Number of Plans", self.num_plans],       #1
                         ["Objective Function", self.objective],    #2
                         ["Planning Horizon", self.horizon],        #3
                         ["Constraints", self.constraints],         #4
                         ["Statistical Life", self.stat_life],      #5
                         ["Discount Rate", self.discount_rate],     #6
                         ["Disaster Rate", self.disaster_rate],     #7
                         ["Disaster Magnitude", self.dis_magnitude],#8
                         ["Risk Preference", self.risk_preference], #9
                         ["-", "Base"],                             #10
                         ["Name of Plans"],                         #11
                         [""],                                      #12
                         ["Costs"],                                 #13
                         ["Direct Costs"],                          #14
                         ["Indirect Costs"],                        #15
                         ["OMR"],                                   #16
                         ["OMRType"],                               #17
                         ["Externalities"],                         #18
                         ["Benefits"],                              #19
                         ["Direct Loss Reduction"],                 #20
                         ["Indirect Loss Reduction"],               #21
                         ["Response/Recovery Costs"],               #22
                         ["Non Disaster Related Benefits"],         #23
                         ["NDRBType"],                              #24
                         ["Fatalities Averted"]]                    #25

        # ===== The following 3 for loops are for Excel spreadsheet neatness
        for i in range(self.num_plans+1):
            if i != 0:
                to_write_list[10].append("Plan " + str(i))

            for j in range((max_lengths[i] * self.FIELDS_PER_VALUE) - 1):
                to_write_list[10].append("")
#        to_write_list[10].append("/")      # === Marks end of list of plans (for file reading)

        for i in range(self.num_plans+1):
            to_write_list[11].append(self.plan_name[i])
            for j in range((max_lengths[i]*self.FIELDS_PER_VALUE) - 1):
                to_write_list[11].append("")

        for i in range(self.num_plans+1):
            for j in range(max_lengths[i]):
                to_write_list[12].extend(["Title", "Dollar Amount", "Description"])

        # ===== Direct Cost
        for i in range(self.num_plans+1):
            for j in range(len(self.cost.direct[i])):
                for k in range(self.FIELDS_PER_VALUE):
                    to_write_list[14].append(self.cost.direct[i][j][k])

            for j in range(max_lengths[i] - len(self.cost.direct[i])):
                to_write_list[14].extend([""]*self.FIELDS_PER_VALUE)

        # ===== Indirect Cost
        for i in range(self.num_plans + 1):
            for j in range(len(self.cost.indirect[i])):
                for k in range(self.FIELDS_PER_VALUE):
                    to_write_list[15].append(self.cost.indirect[i][j][k])

            for j in range(max_lengths[i] - len(self.cost.indirect[i])):
                to_write_list[15].extend([""]*self.FIELDS_PER_VALUE)

        # ===== OMR
        for i in range(self.num_plans + 1):
            for j in range(len(self.cost.omr[i])):
                for k in range(self.FIELDS_PER_VALUE):
                    to_write_list[16].append(self.cost.omr[i][j][k])

            for j in range(max_lengths[i] - len(self.cost.omr[i])):
                to_write_list[16].extend([""]*self.FIELDS_PER_VALUE)

        # ===== OMRType
        for i in range(self.num_plans + 1):
            for j in range(len(self.cost.omr_type[i])):
                for k in range(self.FIELDS_PER_VALUE):
                    to_write_list[17].append(self.cost.omr_type[i][j][k])

            for j in range(max_lengths[i] - len(self.cost.omr_type[i])):
                to_write_list[17].extend([""]*self.FIELDS_PER_VALUE)

        # ===== Externalities
        for i in range(self.num_plans + 1):
            for j in range(len(self.Externalities[i])):
                for k in range(self.FIELDS_PER_VALUE):
                    to_write_list[18].append(self.Externalities[i][j][k])

            for j in range(max_lengths[i] - len(self.Externalities[i])):
                to_write_list[18].extend([""]*self.FIELDS_PER_VALUE)

        # ===== Direct Loss Reductions
        for i in range(self.num_plans + 1):
            for j in range(len(self.ben.direct[i])):
                for k in range(self.FIELDS_PER_VALUE):
                    to_write_list[20].append(self.ben.direct[i][j][k])

            for j in range(max_lengths[i] - len(self.ben.direct[i])):
                to_write_list[20].extend([""]*self.FIELDS_PER_VALUE)

        # ===== Indirect Loss Reductions
        for i in range(self.num_plans + 1):
            for j in range(len(self.ben.indirect[i])):
                for k in range(self.FIELDS_PER_VALUE):
                    to_write_list[21].append(self.ben.indirect[i][j][k])

            for j in range(max_lengths[i] - len(self.ben.indirect[i])):
                to_write_list[21].extend([""]*self.FIELDS_PER_VALUE)

        # ===== Response/Recovery Cost Reductions
        for i in range(self.num_plans + 1):
            for j in range(len(self.ben.res_rec[i])):
                for k in range(self.FIELDS_PER_VALUE):
                    to_write_list[22].append(self.ben.res_rec[i][j][k])

            for j in range(max_lengths[i] - len(self.ben.res_rec[i])):
                to_write_list[22].extend([""]*self.FIELDS_PER_VALUE)

        # ===== Non Disaster Related Benefits
        for i in range(self.num_plans + 1):
            for j in range(len(self.NonDBen[i])):
                for k in range(self.FIELDS_PER_VALUE):
                    to_write_list[23].append(self.NonDBen[i][j][k])

            for j in range(max_lengths[i] - len(self.NonDBen[i])):
                to_write_list[23].extend([""]*self.FIELDS_PER_VALUE)

        # ===== Non Disaster Related Benefits Recursion
        for i in range(self.num_plans + 1):
            for j in range(len(self.NonDBen[i])):
                if len(self.NonDBen[i][j]) < 6:
                    self.NonDBen[i][j].extend([""] * (4 - len(self.NonDBen[j])))
                for k in range(3, len(self.NonDBen[i][j])):
                    to_write_list[24].append(self.NonDBen[i][j][k])

            for j in range(max_lengths[i] - len(self.NonDBen[i])):
                to_write_list[24].extend([""] * self.FIELDS_PER_VALUE)

        # ===== Fatalities Averted
        for i in range(self.num_plans + 1):
            for j in range(self.FIELDS_PER_VALUE):
                to_write_list[25].append(self.Fatalities[i][j])

            for j in range(max_lengths[i] - 1):
                to_write_list[25].extend([""] * self.FIELDS_PER_VALUE)

        # ============ Prompts the user to select a location to save to (csv format)
        my_formats = [
            ('Comma Separated Value', '*.csv'),
        ]
        file_name = filedialog.asksaveasfilename(filetypes=my_formats, title="Save the file as...")
        if '.csv' not in file_name:
            file_name = file_name + '.csv'
        file = open(file_name, 'w')


        # ============ Take to_write_list and save into the user inputted location (as a .csv file)
        for row in range(10):
            for value in range(2):
                file.write(str(to_write_list[row][value]))
                file.write(',')
            file.write('\n')
        count = 0

        for i in range(self.num_plans + 1):
            for j in range(max_lengths[i]):
                count += self.FIELDS_PER_VALUE
        count += 2
        for value in range(len(to_write_list[10])):
            file.write(to_write_list[10][value])
            file.write(',')
        file.write('\n')

        count -= 1
        for row in range(11, 13):
            for value in range(count):
                file.write(str(to_write_list[row][value]))
                file.write(',')
            file.write('\n')
        file.write(str(to_write_list[13][0]))
        file.write('\n')

        for row in range(14, 19):
            for value in range(count):
                file.write(str(to_write_list[row][value]))
                file.write(',')
            file.write('\n')
        file.write(str(to_write_list[19][0]))
        file.write('\n')

        for row in range(20, 26):
            for value in range(count):
                file.write(str(to_write_list[row][value]))
                file.write(',')
            file.write('\n')

    def to_docx(self):
        """
        Allows user to save current analysis as a .docx file (for a complete report in paper-form)
        """
        my_formats = [
            ('Microsoft Word Document', '*.docx'),
        ]
        file_name = filedialog.asksaveasfilename(filetypes=my_formats, title="Save the file as...")
        if '.docx' not in file_name:
            file_name = file_name + '.docx'

        doc = docx.Document()               # === Opens a new .docx document

        doc.add_heading('Economic Evaluation Complete Report\n' + self.analysis_title, 0)

        doc.add_heading('Analysis Base Information\n', 1)
        doc.add_paragraph(TAB+'Number of Alternatives: '+str(self.num_plans))
        doc.add_paragraph(TAB+'Planning Horizon: '+str(self.horizon)+' years')
        doc.add_paragraph(TAB+'Discount Rate: '+str(self.discount_rate)+'%')
        doc.add_paragraph()
        doc.add_paragraph(TAB+'Disaster Rate: Every '+str(self.disaster_rate)+' years')
        doc.add_paragraph(TAB+'Disaster Magnitude: '+str(self.dis_magnitude)+'% of build cost')
        doc.add_paragraph(TAB+'Risk Preference: '+str(self.risk_preference))
        doc.add_paragraph()
        doc.add_paragraph(TAB+'Statistical Life Value: '+'${:,.0f}'.format(float(self.stat_life)))

        for i in range(self.num_plans + 1):
            if i == 0:
                doc.add_heading('Base Plan', 1)
            else:
                doc.add_heading(self.plan_name[i] + " (Alternative " + str(i) + ")", 1)

            doc.add_heading('Disaster-Related Benefits\n', 2)
            # === Makes sure the field isn't blank
            if self.ben.res_rec[i][0][0] != "":
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(TAB + 'Response and Recovery Cost Reductions\n')
                # === Allows the text to be bolded
                run.bold = True
                for j in range(len(self.ben.res_rec[i])):
                    doc.add_paragraph(TAB + str(j+1) + ") " + self.ben.res_rec[i][j][0])
                    dollar_amount = str('${:,.0f}'.format(float(self.ben.res_rec[i][j][1])))
                    doc.add_paragraph(TAB + 'Dollar Amount: ' + dollar_amount)
                    present_value = self.on_dis_occ(float(self.ben.res_rec[i][j][1]))
                    form_present_value = '${:,.0f}'.format(present_value)
                    doc.add_paragraph(TAB + 'Effective Present Value: ' + form_present_value)
                    doc.add_paragraph(TAB + 'Description: ' + str(self.ben.res_rec[i][j][2]) + "\n")

            if self.ben.direct[i][0][0] != "":
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(TAB + 'Direct Losses Prevented\n')
                run.bold = True
                for j in range(len(self.ben.direct[i])):
                    doc.add_paragraph(TAB + str(j+1) + ") " + self.ben.direct[i][j][0])
                    dollar_amount = str('${:,.0f}'.format(float(self.ben.direct[i][j][1])))
                    doc.add_paragraph(TAB + 'Dollar Amount: ' + dollar_amount)
                    present_value = self.on_dis_occ(float(self.ben.direct[i][j][1]))
                    form_present_value = '${:,.0f}'.format(present_value)
                    doc.add_paragraph(TAB + 'Effective Present Value: ' + form_present_value)
                    doc.add_paragraph(TAB + 'Description: ' + str(self.ben.direct[i][j][2]) + "\n")

            if self.ben.indirect[i][0][0] != "":
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(TAB + 'Indirect Losses Prevented\n')
                run.bold = True
                for j in range(len(self.ben.indirect[i])):
                    doc.add_paragraph(TAB + str(j + 1) + ") " + self.ben.indirect[i][j][0])
                    doc.add_paragraph(TAB + 'Dollar Amount: '
                                      + str('${:,.0f}'.format(float(self.ben.indirect[i][j][1]))))
                    new_text = self.on_dis_occ(float(self.ben.indirect[i][j][1]))
                    doc.add_paragraph(TAB + 'Effective Present Value: '
                                      + '${:,.0f}'.format(new_text))
                    doc.add_paragraph(TAB + 'Description: '
                                      + str(self.ben.indirect[i][j][2]) + "\n")

            if float(self.Fatalities[i][1]) != 0:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(TAB + 'Fatalities Averted\n')
                run.bold = True
                doc.add_paragraph(TAB + 'Number of Fatalities Averted: '
                                  + str(self.Fatalities[i][1]))
                doc.add_paragraph(TAB + 'Dollar Worth: '
                                  + str('${:,.0f}'.format(self.fatalities_sum[i])))
                doc.add_paragraph(TAB + 'Description: ' + str(self.Fatalities[i][2]) + "\n")

            if self.NonDBen[i][0][0] != "":
                doc.add_heading('Non-disaster-Related Benefits (Resilience Dividend)\n', 2)
                for j in range(len(self.NonDBen[i])):
                    if self.NonDBen[i][j][3] == "OneTime":
                        doc.add_paragraph(TAB + str(j+1) + ") " + self. NonDBen[i][j][0])
                        doc.add_paragraph(TAB + 'Dollar Amount: '
                                          + str('${:,.0f}'.format(float(self.NonDBen[i][j][1]))))
                        doc.add_paragraph(TAB + 'Year Applied: '
                                          + str(self.NonDBen[i][j][4]
                                                + ' years after plan start date'))
                        new_text = self.calc_one_time(float(self.NonDBen[i][j][1]),
                                                      float(self.NonDBen[i][j][4]))
                        doc.add_paragraph(TAB + 'Effective Present Value: '
                                          + '${:,.0f}'.format(new_text))
                        doc.add_paragraph(TAB + 'Description: ' + str(self.NonDBen[i][j][2]) + "\n")
                    if self.NonDBen[i][j][3] == "Recurring":
                        doc.add_paragraph(TAB + str(j + 1) + ") " + self.NonDBen[i][j][0])
                        doc.add_paragraph(TAB + 'Dollar Amount: '
                                          + str('${:,.0f}'.format(float(self.NonDBen[i][j][1]))))
                        doc.add_paragraph(TAB + 'Years Applied: '
                                          + str(self.NonDBen[i][j][4])
                                          + ' years after plan start date and every '
                                          + str(self.NonDBen[i][j][5]) + ' year(s) afterwards')
                        new_text = self.calc_recur(float(self.NonDBen[i][j][1]),
                                                   float(self.NonDBen[i][j][4]),
                                                   float(self.NonDBen[i][j][5]))
                        doc.add_paragraph(TAB + 'Effective Present Value: '
                                          + '${:,.0f}'.format(new_text))
                        doc.add_paragraph(TAB + 'Description: ' + str(self.NonDBen[i][j][2]) + "\n")

            doc.add_heading('Costs\n', 2)
            if self.cost.direct[i][0][0] != "":
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(TAB + 'Up-Front Direct Costs\n')
                run.bold = True
                for j in range(len(self.cost.direct[i])):
                    doc.add_paragraph(TAB + str(j + 1) + ") " + self.cost.direct[i][j][0])
                    doc.add_paragraph(TAB + 'Dollar Amount: '
                                      + str('${:,.0f}'.format(float(self.cost.direct[i][j][1]))))
                    doc.add_paragraph(TAB + 'Effective Present Value: '
                                      + '${:,.0f}'.format(float(self.cost.direct[i][j][1])))
                    doc.add_paragraph(TAB + 'Description: ' + str(self.cost.direct[i][j][2]) + "\n")

            if self.cost.indirect[i][0][0] != "":
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(TAB + 'Up-Front Indirect Costs\n')
                run.bold = True
                for j in range(len(self.cost.indirect[i])):
                    doc.add_paragraph(TAB + str(j + 1) + ") " + self.cost.indirect[i][j][0])
                    doc.add_paragraph(TAB + 'Dollar Amount: '
                                      + str('${:,.0f}'.format(float(self.cost.indirect[i][j][1]))))
                    doc.add_paragraph(TAB + 'Effective Present Value: '
                                      + '${:,.0f}'.format(float(self.cost.indirect[i][j][1])))
                    doc.add_paragraph(TAB + 'Description: '
                                      + str(self.cost.indirect[i][j][2]) + "\n")

            has_one_time = False
            for j in range(len(self.cost.omr_type[i])):
                # === Checks if there are ANY One-time OMRs to report
                if self.cost.omr_type[i][j][0] == "OneTime":
                    has_one_time = True
            if has_one_time:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(TAB + 'One-time Operation, Management, or Repairs Costs\n')
                run.bold = True
                for j in range(len(self.cost.omr[i])):
                    if self.cost.omr_type[i][j][0] == "OneTime":
                        doc.add_paragraph(TAB + str(j + 1) + ") " + self.cost.omr[i][j][0])
                        doc.add_paragraph(TAB + 'Dollar Amount: '
                                          + str('${:,.0f}'.format(float(self.cost.omr[i][j][1]))))
                        doc.add_paragraph(TAB + 'Year Applied: '
                                          + str(self.cost.omr_type[i][j][1]
                                                + ' years after plan start date'))
                        new_text = self.calc_one_time(float(self.cost.omr[i][j][1]),
                                                      float(self.cost.omr_type[i][j][1]))
                        doc.add_paragraph(TAB + 'Effective Present Value: '
                                          + '${:,.0f}'.format(new_text))
                        doc.add_paragraph(TAB + 'Description: '
                                          + str(self.cost.omr[i][j][2]) + "\n")

            has_recurring = False
            for j in range(len(self.cost.omr_type[i])):
                # === Checks if there are ANY Recurring OMRs to report
                if self.cost.omr_type[i][j][0] == "Recurring":
                    has_recurring = True
            if has_recurring:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(TAB + 'Recurring Operation, Management, or Repairs Costs\n')
                run.bold = True
                for j in range(len(self.cost.omr[i])):
                    if self.cost.omr_type[i][j][0] == "Recurring":
                        doc.add_paragraph(TAB + str(j + 1) + ") " + self.cost.omr[i][j][0])
                        doc.add_paragraph(TAB + 'Dollar Amount: '
                                          + str('${:,.0f}'.format(float(self.cost.omr[i][j][1]))))
                        doc.add_paragraph(TAB + 'Years Applied: '
                                          + str(self.cost.omr_type[i][j][1])
                                          + ' years after plan start date'
                                          ' and every ' + str(self.cost.omr_type[i][j][2])
                                          + ' year(s) afterwards')
                        new_text = self.calc_recur(float(self.cost.omr[i][j][1]),
                                                   float(self.cost.omr_type[i][j][1]),
                                                   float(self.cost.omr_type[i][j][2]))
                        doc.add_paragraph(TAB + 'Effective Present Value: '
                                          + '${:,.0f}'.format(new_text))
                        doc.add_paragraph(TAB + 'Description: '
                                          + str(self.cost.omr[i][j][2]) + "\n")

            doc.add_heading('Totals\n', 2)

            paragraph = doc.add_paragraph()
            if self.net[i] >= 0:
                run = paragraph.add_run(TAB + 'Total Benefits: '
                                        + '${:,.0f}'.format(self.ben.total[i])
                                        + '\n' + TAB + 'Total Costs: '
                                        + '${:,.0f}'.format(self.cost.total[i])
                                        + '\n' + TAB + 'Net: '
                                        + '${:,.0f}'.format(self.net[i]))
            else:
                run = paragraph.add_run(TAB + 'Total Benefits: '
                                        + '${:,.0f}'.format(self.ben.total[i])
                                        + '\n' + TAB + 'Total Costs: '
                                        + '${:,.0f}'.format(self.cost.total[i])
                                        + '\n' + TAB + 'Net: ('
                                        + '${:,.0f}'.format(self.net[i])+')')
            run.bold = True
            doc.add_paragraph()
            doc.add_paragraph(TAB + 'Savings-to-Investment Ratio (SIR): '
                              + '{:,.2f}'.format(self.sir(i)))
            try:
                text = '{:,.1f}'.format(self.irr(i))+'%'
            except ValueError:
                text = self.irr(i)
            doc.add_paragraph(TAB+'Internal Rate of Return (IRR): '+ text)
            doc.add_paragraph(TAB+'Return on Investment (ROI): '+'{:,.1f}'.format(self.roi(i))+'%')
            doc.add_paragraph(TAB+'Non-Disaster ROI: '+'{:,.1f}'.format(self.non_d_roi(i))+'%')

            doc.save(file_name)


    def to_csv(self):
        """Allows the user to save summary of analysis as a .csv file
        (not to be confused with the to_save .csv file)"""
        to_write_list = [["Outputs of Economic Evaluation: [" + str(self.analysis_title) + "]"], #0
                         [" ", "Base Case"],                                                     #1
                         [" "],                                                                  #2
                         ["Benefits"],                                                           #3
                         ["Disaster Economic Benefits"],                                         #4
                         ["Response and Recovery Costs"],                                        #5
                         ["Direct Losses Prevented"],                                            #6
                         ["Indirect Losses Prevented"],                                          #7
                         ["Disaster Non-Market Benefits"],                                       #8
                         ["Lives Saved"],                                                        #9
                         ["Non-disaster Related Benefits"]]                                      #10

        to_write_list.extend([["Costs"],                                              #11
                              ["Initial"],                                            #12
                              ["Direct Cost"],                                        #13
                              ["Indirect Cost"],                                      #14
                              ["OMR"],                                                #15
                              ["Recurring Costs"],                                    #16
                              ["OMR"],                                                #17
                              ["Total: Present Expected Value"],                      #18
                              ["Benefits"],                                           #19
                              ["Costs"],                                              #20
                              ["Net"],                                                #21
                              [" "],                                                  #22
                              ["Savings-to-Investment Ratio (SIR)"],                  #23
                              ["Internal Rate of Return (IRR)"],                      #24
                              ["Return on Investment (ROI)"],                         #25
                              ["Non-Disaster ROI"]])                                  #26

        for i in range(1, self.num_plans + 1):
            to_write_list[1].append("Alternative " + str(i))
        for i in range(self.num_plans + 1):
            to_write_list[0].append(" ")
            to_write_list[2].append(self.plan_name[i])
            to_write_list[3].append(" ")
            to_write_list[4].append(" ")
            to_write_list[5].append('${:.0f}'.format(self.on_dis_occ(self.ben.r_sum[i])))
            to_write_list[6].append('${:.0f}'.format(self.on_dis_occ(self.ben.d_sum[i])))
            to_write_list[7].append('${:.0f}'.format(self.on_dis_occ(self.ben.i_sum[i])))
            to_write_list[8].append(" ")
            to_write_list[9].append('${:.0f}'.format(self.calc_fatalities(self.Fatalities[i][1])))
            # == Calculation NDRB sum
            total = 0
            for j in range(len(self.NonDBen[i])):
                if self.NonDBen[i][j][3] == 'OneTime':
                    total += self.calc_one_time(self.NonDBen[i][j][1], self.NonDBen[i][j][4])
                elif self.NonDBen[i][j][3] == 'Recurring':
                    total += self.calc_recur(self.NonDBen[i][j][1],
                                             self.NonDBen[i][j][4], self.NonDBen[i][j][5])
            to_write_list[10].append('${:.0f}'.format(total))

        for i in range(self.num_plans + 1):
            to_write_list[11].append(" ")
            to_write_list[12].append(" ")
            to_write_list[13].append('${:.0f}'.format(self.cost.d_sum[i]))
            to_write_list[14].append('${:.0f}'.format(self.cost.i_sum[i]))
            to_write_list[15].append('${:.0f}'.format(self.cost.omr_1_sum[i]))
            to_write_list[16].append(" ")
            to_write_list[17].append('${:.0f}'.format(self.cost.omr_r_sum[i]))
            to_write_list[18].append(" ")
            to_write_list[19].append('${:.0f}'.format(self.ben.total[i]))
            to_write_list[20].append('${:.0f}'.format(self.cost.total[i]))
            if self.net[i] >= 0:
                to_write_list[21].append('${:.0f}'.format(self.net[i]))
            else:
                to_write_list[21].append('(${:.0f}'.format(self.net[i]) + ')')
            to_write_list[22].append(" ")
            to_write_list[23].append('{:.2f}'.format(self.sir(i)))
            if type(self.irr(i)) == type("string"):
                to_write_list[24].append(self.irr(i))
            else:
                to_write_list[24].append('{:.1f}'.format(self.irr(i)) + '%')
            to_write_list[25].append('{:.1f}'.format(self.roi(i)) + '%')
            to_write_list[26].append('{:.1f}'.format(self.non_d_roi(i)) + '%')


        # ============ Prompts the user to select a location to save to (csv format)
        my_formats = [('Comma Separated Value', '*.csv'),]
        file_name = filedialog.asksaveasfilename(filetypes=my_formats, title="Save the file as...")
        if '.csv' not in file_name:
            file_name = file_name + '.csv'
        file = open(file_name, 'w')

        # ============ Take to_write_list and save into the user inputted location (as a .csv file)
        for row in range(26):
            for value in range(self.num_plans + 2):
                file.write(str(to_write_list[row][value]))
                file.write(',')
            file.write('\n')
